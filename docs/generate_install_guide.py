"""cmux-win 설치 가이드 Word 문서 생성"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

doc = Document()

# ── 페이지 설정 ──
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

# ── 스타일 설정 ──
style = doc.styles['Normal']
font = style.font
font.name = 'Pretendard'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.4

# 한글 폰트 fallback
rFonts = style.element.find(qn('w:rPr'))
if rFonts is None:
    rPr = style.element.makeelement(qn('w:rPr'), {})
    style.element.append(rPr)
    rFonts = rPr

for heading_level in range(1, 4):
    h_style = doc.styles[f'Heading {heading_level}']
    h_style.font.name = 'Pretendard'
    h_style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def add_para(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_code_block(text):
    """코드 블록 (회색 배경, 모노스페이스)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # 배경색은 shading으로
    shading = p.paragraph_format.element.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F0F2F5',
        qn('w:val'): 'clear',
    })
    pPr = p.paragraph_format.element
    if pPr.find(qn('w:pPr')) is None:
        pPr_elem = p.paragraph_format.element.makeelement(qn('w:pPr'), {})
        pPr.insert(0, pPr_elem)
    p._element.find(qn('w:pPr')).append(shading)

    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return p


def add_step(number, title, description=''):
    """번호 + 제목 + 설명"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run_num = p.add_run(f'Step {number}. ')
    run_num.bold = True
    run_num.font.size = Pt(12)
    run_num.font.color.rgb = RGBColor(0x00, 0x91, 0xFF)
    run_title = p.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(12)
    if description:
        p2 = doc.add_paragraph(description)
        p2.paragraph_format.left_indent = Cm(0.5)
    return p


def add_tip(text):
    """팁 박스"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run_icon = p.add_run('💡 ')
    run_icon.font.size = Pt(11)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return p


def add_warning(text):
    """경고 박스"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run_icon = p.add_run('⚠️ ')
    run_icon.font.size = Pt(11)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x55, 0x00)
    return p


# ════════════════════════════════════════════════════════════════════════
# 문서 본문 시작
# ════════════════════════════════════════════════════════════════════════

# ── 표지 ──
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('cmux-win')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x00, 0x91, 0xFF)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('설치 및 실행 가이드')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

doc.add_paragraph()
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('Windows AI Terminal Multiplexer')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.italic = True

desc2 = doc.add_paragraph()
desc2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc2.add_run('여러 AI CLI(Claude, AGY, Codex)를 동시에 실행하고 협업시키는 데스크톱 앱')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ver.add_run('v0.1.0  |  2026년 6월')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════
# 1. 사전 준비
# ══════════════════════════════════════════════════════════════════════
add_heading('1. 사전 준비', level=1)

add_para('cmux-win을 설치하기 전에 아래 프로그램들이 필요합니다.')

# 테이블: 사전 준비 항목
table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['프로그램', '버전', '다운로드']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

data = [
    ['Git', '2.40 이상', 'https://git-scm.com/downloads'],
    ['Node.js', 'v20 이상 (LTS 권장)', 'https://nodejs.org'],
    ['Python', '3.10 이상', 'https://python.org (선택)'],
    ['Visual Studio\nBuild Tools', '2019 이상', 'npm install 시 자동 설치 시도'],
]
for row_idx, row_data in enumerate(data, start=1):
    for col_idx, val in enumerate(row_data):
        table.rows[row_idx].cells[col_idx].text = val

doc.add_paragraph()
add_tip('Node.js 설치 시 "Add to PATH" 옵션을 반드시 체크하세요.')
add_tip('Python은 Streamlit 대시보드를 사용할 때만 필요합니다.')

doc.add_paragraph()
add_heading('설치 확인 방법', level=2)
add_para('명령 프롬프트(cmd) 또는 PowerShell을 열고 아래 명령을 실행합니다:')
add_code_block('git --version')
add_code_block('node --version')
add_code_block('npm --version')
add_para('각각 버전 번호가 표시되면 준비 완료입니다.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════
# 2. 소스 코드 다운로드 (git clone)
# ══════════════════════════════════════════════════════════════════════
add_heading('2. 소스 코드 다운로드', level=1)

add_step(1, 'GitHub에서 소스 코드 복제',
         '원하는 폴더에서 아래 명령을 실행합니다.')
add_code_block('git clone https://github.com/yijae78/cmux-win.git')

add_step(2, '프로젝트 폴더로 이동')
add_code_block('cd cmux-win')

add_tip('기본적으로 현재 폴더 안에 cmux-win 폴더가 생성됩니다.')
add_tip('예: C:\\dev 에서 실행하면 C:\\dev\\cmux-win 이 생성됩니다.')

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════
# 3. 의존성 설치 (npm install)
# ══════════════════════════════════════════════════════════════════════
add_heading('3. 의존성 설치', level=1)

add_step(1, 'npm install 실행',
         '프로젝트의 모든 라이브러리를 자동으로 다운로드합니다.')
add_code_block('npm install')

add_para('설치에 3~5분 정도 소요됩니다. 특히 node-pty와 better-sqlite3는 C++ 네이티브 모듈이므로 빌드 도구가 필요합니다.')

add_warning('만약 node-gyp 관련 에러가 발생하면, Visual Studio Build Tools를 설치해야 합니다.')

add_para('Build Tools 자동 설치 (관리자 권한 PowerShell에서):')
add_code_block('npm install --global windows-build-tools')

add_para('또는 Visual Studio Installer에서 "C++ 빌드 도구" 워크로드를 선택하여 설치합니다.')

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════
# 4. 빌드
# ══════════════════════════════════════════════════════════════════════
add_heading('4. 빌드', level=1)

add_step(1, '앱 빌드',
         'TypeScript 소스를 JavaScript로 컴파일합니다.')
add_code_block('npm run build')

add_para('빌드가 성공하면 out/ 폴더에 결과물이 생성됩니다:')

table2 = doc.add_table(rows=4, cols=2)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
table2.rows[0].cells[0].text = '폴더'
table2.rows[0].cells[1].text = '설명'
for p in table2.rows[0].cells[0].paragraphs:
    for r in p.runs: r.bold = True
for p in table2.rows[0].cells[1].paragraphs:
    for r in p.runs: r.bold = True
build_data = [
    ['out/main/', 'Electron 메인 프로세스 (백엔드)'],
    ['out/preload/', 'IPC 브릿지 (보안 계층)'],
    ['out/renderer/', 'React UI (프론트엔드)'],
]
for i, (folder, desc_text) in enumerate(build_data, 1):
    table2.rows[i].cells[0].text = folder
    table2.rows[i].cells[1].text = desc_text

doc.add_paragraph()

add_step(2, 'MCP 서버 빌드 (선택)',
         'Claude Desktop에서 cmux-win을 제어하려면 MCP 서버도 빌드합니다.')
add_code_block('npm run build:mcp')

add_tip('MCP 서버는 Claude Desktop의 도구(tool)로 등록되어, 핸드폰에서도 씨윈을 제어할 수 있게 합니다.')

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════
# 5. 실행
# ══════════════════════════════════════════════════════════════════════
add_heading('5. 실행', level=1)

add_step(1, '앱 실행')
add_code_block('npx electron out/main/index.js')

add_para('앱이 실행되면:')
p = doc.add_paragraph(style='List Bullet')
p.add_run('왼쪽 사이드바에 워크스페이스 목록이 표시됩니다')
p = doc.add_paragraph(style='List Bullet')
p.add_run('첫 번째 터미널에 Claude CLI가 자동으로 실행됩니다')
p = doc.add_paragraph(style='List Bullet')
p.add_run('소켓 API가 localhost:19840에서 대기합니다')

doc.add_paragraph()
add_tip('개발 모드(핫 리로드)로 실행하려면: npm run dev')

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════
# 6. 바탕화면 바로가기 만들기 (선택)
# ══════════════════════════════════════════════════════════════════════
add_heading('6. 바탕화면 바로가기 만들기 (선택)', level=1)

add_para('매번 명령어를 입력하지 않으려면 바탕화면 바로가기를 만들 수 있습니다.')
add_para('PowerShell에서 아래 명령을 실행합니다 (경로를 본인 환경에 맞게 수정):')

shortcut_code = '''$project = "C:\\dev\\cmux-win"
$electron = "$project\\node_modules\\electron\\dist\\electron.exe"
$s = (New-Object -COM WScript.Shell).CreateShortcut(
    "$env:USERPROFILE\\Desktop\\cmux-win.lnk")
$s.TargetPath = $electron
$s.Arguments = "out/main/index.js"
$s.WorkingDirectory = $project
$s.IconLocation = "$electron,0"
$s.Save()'''
add_code_block(shortcut_code)

add_tip('바탕화면에 cmux-win.lnk 아이콘이 생성됩니다. 더블클릭으로 실행하세요.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════
# 7. 패키지 빌드 (설치 파일 만들기)
# ══════════════════════════════════════════════════════════════════════
add_heading('7. 설치 파일 만들기 (선택)', level=1)

add_para('.exe 설치 파일을 만들어서 다른 PC에 배포할 수 있습니다:')
add_code_block('npm run package')

add_para('빌드가 완료되면 dist/ 폴더에 설치 파일이 생성됩니다.')

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════
# 8. 주요 기능 사용법
# ══════════════════════════════════════════════════════════════════════
add_heading('8. 주요 기능 간단 안내', level=1)

add_heading('터미널 패널', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('패널 분할: ').bold = True
p.add_run('터미널 상단 바의 [+] 버튼 또는 Ctrl+Shift+\\')
p = doc.add_paragraph(style='List Bullet')
p.add_run('균등 분할: ').bold = True
p.add_run('사이드바의 ⊞ (가로) / ⊟ (세로) 버튼')
p = doc.add_paragraph(style='List Bullet')
p.add_run('탭 전환: ').bold = True
p.add_run('패널 상단의 탭을 클릭')

add_heading('AI CLI 실행', level=2)

table3 = doc.add_table(rows=4, cols=2)
table3.style = 'Light Grid Accent 1'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
table3.rows[0].cells[0].text = 'AI'
table3.rows[0].cells[1].text = '실행 명령'
for p in table3.rows[0].cells[0].paragraphs:
    for r in p.runs: r.bold = True
for p in table3.rows[0].cells[1].paragraphs:
    for r in p.runs: r.bold = True
cli_data = [
    ['Claude', 'claude  (자동 실행됨)'],
    ['AGY', 'agy -i "작업내용" -y'],
    ['Codex', 'codex --full-auto --no-alt-screen "작업"'],
]
for i, (ai, cmd) in enumerate(cli_data, 1):
    table3.rows[i].cells[0].text = ai
    table3.rows[i].cells[1].text = cmd

doc.add_paragraph()

add_heading('워크스페이스', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('새 워크스페이스: ').bold = True
p.add_run('사이드바 상단의 [+] 버튼')
p = doc.add_paragraph(style='List Bullet')
p.add_run('전환: ').bold = True
p.add_run('사이드바에서 워크스페이스 이름 클릭')
p = doc.add_paragraph(style='List Bullet')
p.add_run('순서 변경: ').bold = True
p.add_run('왼쪽 6점 핸들(⠿)을 잡고 위/아래로 드래그')
p = doc.add_paragraph(style='List Bullet')
p.add_run('이름 변경: ').bold = True
p.add_run('워크스페이스 이름을 더블클릭')

add_heading('파일 탐색기', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('토글: ').bold = True
p.add_run('Ctrl+E 또는 사이드바 하단 ⌂ 아이콘')
p = doc.add_paragraph(style='List Bullet')
p.add_run('폴더 열기: ').bold = True
p.add_run('탐색기에서 폴더를 더블클릭하면 터미널에 cd 전송')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════
# 9. 문제 해결
# ══════════════════════════════════════════════════════════════════════
add_heading('9. 문제 해결', level=1)

# 테이블: 문제-해결
table4 = doc.add_table(rows=6, cols=2)
table4.style = 'Light Grid Accent 1'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
table4.rows[0].cells[0].text = '문제'
table4.rows[0].cells[1].text = '해결 방법'
for p in table4.rows[0].cells[0].paragraphs:
    for r in p.runs: r.bold = True
for p in table4.rows[0].cells[1].paragraphs:
    for r in p.runs: r.bold = True

trouble_data = [
    ['npm install 에서\nnode-gyp 에러',
     'Visual Studio Build Tools 설치 후 재시도\n'
     '또는: npm install --global windows-build-tools'],
    ['electron 실행 시\n하얀 화면만 표시',
     'npm run build 를 먼저 실행했는지 확인\n'
     'out/ 폴더가 존재하는지 확인'],
    ['Claude CLI가\n자동 실행되지 않음',
     'claude 명령이 PATH에 있는지 확인:\n'
     'claude --version'],
    ['소켓 연결 실패\n(포트 19840)',
     '다른 프로그램이 포트를 사용 중인지 확인:\n'
     'netstat -an | findstr 19840'],
    ['앱이 2개 실행됨',
     '싱글 인스턴스 잠금이 작동합니다.\n'
     '기존 앱을 종료하고 다시 실행하세요.'],
]
for i, (problem, solution) in enumerate(trouble_data, 1):
    table4.rows[i].cells[0].text = problem
    table4.rows[i].cells[1].text = solution


doc.add_paragraph()
doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════
# 10. 전체 명령어 요약
# ══════════════════════════════════════════════════════════════════════
add_heading('10. 전체 명령어 요약 (한눈에 보기)', level=1)

add_para('처음부터 끝까지, 아래 명령을 순서대로 실행하면 됩니다:')

summary = '''# 1. 소스 코드 다운로드
git clone https://github.com/yijae78/cmux-win.git
cd cmux-win

# 2. 의존성 설치
npm install

# 3. 빌드
npm run build

# 4. 실행
npx electron out/main/index.js'''
add_code_block(summary)

add_para('이것이 전부입니다!', bold=True)

doc.add_paragraph()
doc.add_paragraph()

# 푸터
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('— cmux-win Installation Guide —')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
run.italic = True


# ── 저장 ──
output_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(output_dir, 'cmux-win_설치가이드.docx')
doc.save(output_path)
print(f'Done: {output_path}')
