"""
카카오뱅크 AI 기반 이상거래탐지(FDS) 고도화 보고서 생성 스크립트
python-docx 기반, 전문 보고서 수준 문서 생성
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# ── Constants ──
KAKAO_YELLOW = RGBColor(0xFF, 0xCD, 0x00)
KAKAO_DARK = RGBColor(0x3A, 0x1D, 0x1D)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MID_GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
HEADER_BG_HEX = "FFCD00"
HEADER_TEXT_COLOR = RGBColor(0x3A, 0x1D, 0x1D)
TABLE_BORDER_HEX = "999999"
TABLE_ALT_ROW_HEX = "FFF8E1"

OUTPUT_DIR = r"C:\dev\cmux-win\outputs"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "KakaoBank_AI_FDS_Report.docx")
FONT_NAME = "맑은 고딕"


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs: top, bottom, left, right, insideH, insideV with (sz, color, val)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        el = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", "4")}" w:space="0" '
            f'w:color="{attrs.get("color", "000000")}"/>'
        )
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_table_borders(table, color="999999", sz="4"):
    """Apply uniform borders to entire table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def set_run_font(run, font_name=FONT_NAME, size=None, bold=False, color=None, italic=False):
    """Configure a run's font properties."""
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_formatted_paragraph(doc, text, font_size=10, bold=False, color=DARK_GRAY,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0,
                            space_after=6, line_spacing=1.15, first_line_indent=None,
                            italic=False):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    set_run_font(run, size=font_size, bold=bold, color=color, italic=italic)
    return p


def add_heading_styled(doc, text, level=1):
    """Add a heading with Malgun Gothic font."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, font_name=FONT_NAME, bold=True,
                     color=KAKAO_DARK if level <= 2 else DARK_GRAY)
    return h


def add_kakao_table(doc, headers, rows, col_widths=None):
    """Add a professional table with KakaoBank yellow header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # Header row
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_BG_HEX)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        set_run_font(run, size=9, bold=True, color=HEADER_TEXT_COLOR)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_ROW_HEX)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            set_run_font(run, size=8.5, color=DARK_GRAY)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    return table


def add_page_break(doc):
    doc.add_page_break()


def add_bullet_point(doc, text, level=0, bold_prefix=None):
    """Add a bullet-pointed paragraph."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        set_run_font(run_b, size=10, bold=True, color=DARK_GRAY)
        run_t = p.add_run(" " + text)
        set_run_font(run_t, size=10, color=DARK_GRAY)
    else:
        run = p.add_run(text)
        set_run_font(run, size=10, color=DARK_GRAY)
    return p


def add_yellow_bar(doc, height_pt=4):
    """Add a KakaoBank yellow divider bar."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    # Use a table with single cell as a bar
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_shading(cell, HEADER_BG_HEX)
    cell.paragraphs[0].text = ""
    cell.height = Pt(height_pt)
    # Set width to full
    cell.width = Cm(16)
    return tbl


def build_cover_page(doc):
    """Build the cover page."""
    # Top spacing
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Yellow top bar
    add_yellow_bar(doc, height_pt=6)

    # Spacing
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("카카오뱅크 AI 기반\n이상거래탐지(FDS) 고도화 보고서")
    set_run_font(run, size=28, bold=True, color=KAKAO_DARK)
    p.paragraph_format.space_after = Pt(12)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("경영전략 및 2026-2030 미래 시나리오")
    set_run_font(run, size=16, color=MID_GRAY)
    p.paragraph_format.space_after = Pt(8)

    # Yellow middle bar
    add_yellow_bar(doc, height_pt=3)

    # Spacing
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Company
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("주식회사 카카오뱅크")
    set_run_font(run, size=14, bold=True, color=DARK_GRAY)
    p.paragraph_format.space_after = Pt(6)

    # Author
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("신교수 · AI 융합학과")
    set_run_font(run, size=12, color=MID_GRAY)
    p.paragraph_format.space_after = Pt(6)

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2026년 4월 29일")
    set_run_font(run, size=12, color=MID_GRAY)
    p.paragraph_format.space_after = Pt(20)

    # Confidential notice
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    set_run_font(run, size=8, color=KAKAO_YELLOW)
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONFIDENTIAL")
    set_run_font(run, size=11, bold=True, color=RGBColor(0xCC, 0x00, 0x00))
    p.paragraph_format.space_after = Pt(3)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("본 보고서는 카카오뱅크 내부 전략 문서로서, 사전 허가 없이\n외부 유출 및 복제를 엄격히 금지합니다.")
    set_run_font(run, size=9, color=MID_GRAY, italic=True)
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    set_run_font(run, size=8, color=KAKAO_YELLOW)


def build_toc(doc):
    """Build Table of Contents page."""
    add_page_break(doc)
    add_heading_styled(doc, "목 차", level=1)

    toc_items = [
        ("Executive Summary", "핵심 요약 및 전략적 결론", "3"),
        ("Chapter 1", "FDS 기술 SOTA 동향", "4"),
        ("Chapter 2", "경쟁사 분석", "6"),
        ("Chapter 3", "규제 동향", "8"),
        ("Chapter 4", "비즈니스 전략 — FDS 투자 ROI 분석", "9"),
        ("Chapter 5", "미래 시나리오 2026-2030", "11"),
        ("Chapter 6", "구현 로드맵", "13"),
        ("Chapter 7", "MLOps & 인프라", "15"),
        ("Chapter 8", "결론 및 전략적 제언", "16"),
        ("Appendix", "용어 해설 및 참고문헌", "18"),
    ]

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    for section, title, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.5
        # Section label
        run_s = p.add_run(f"{section}    ")
        set_run_font(run_s, size=10, bold=True, color=KAKAO_DARK)
        # Title
        run_t = p.add_run(title)
        set_run_font(run_t, size=10, color=DARK_GRAY)
        # Page dots + number
        dots = "  " + "·" * (50 - len(title)) + "  "
        run_d = p.add_run(dots)
        set_run_font(run_d, size=10, color=MID_GRAY)
        run_p = p.add_run(page)
        set_run_font(run_p, size=10, bold=True, color=DARK_GRAY)


def build_executive_summary(doc):
    """Build Executive Summary."""
    add_page_break(doc)
    add_heading_styled(doc, "Executive Summary", level=1)
    add_heading_styled(doc, "핵심 요약 및 전략적 결론", level=2)

    add_formatted_paragraph(
        doc,
        "본 보고서는 카카오뱅크의 AI 기반 이상거래탐지시스템(FDS)의 현황을 분석하고, "
        "2026-2030년 금융 사기 위협 시나리오에 대응하기 위한 기술적·전략적 고도화 방안을 제시합니다. "
        "핵심 기술로 GNN(Graph Neural Network) + Federated Learning 하이브리드 모델을 권고하며, "
        "이를 통해 현재 연간 200억 원의 이상거래 피해액을 50억 원 이하로 감축하는 것을 목표로 합니다.",
        font_size=10, space_before=6, space_after=8
    )

    add_heading_styled(doc, "핵심 결론 3가지", level=3)

    add_bullet_point(doc, "GNN 기반 '금융 지형도'를 구축하여 대포통장 네트워크를 실시간 가시화하고, "
                     "Federated Learning으로 은행 간 사기 패턴을 프라이버시를 보존하면서 공유합니다.",
                     bold_prefix="[기술 혁신]")
    add_bullet_point(doc, "연간 200억 원 사기 피해를 50억 원으로 감축(75% 절감), "
                     "3년 누적 ROI 450억 원 달성을 목표로 합니다. FDS 고도화 투자 총액은 약 150억 원으로 산정됩니다.",
                     bold_prefix="[투자 수익]")
    add_bullet_point(doc, "딥페이크 합성ID 공격(85%), AI 대 AI 적대적 공방(60%), "
                     "CBDC 디지털 자산 연계 사기(75%) 등 3대 미래 위협에 대한 단계별 대응 로드맵을 수립합니다.",
                     bold_prefix="[미래 대응]")

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Summary table
    headers = ["구분", "현재 (2026)", "목표 (2028)", "비전 (2030)"]
    rows = [
        ["연간 사기 피해액", "200억 원", "100억 원", "50억 원"],
        ["FDS 탐지율 (F1-Score)", "0.76", "0.92", "0.97"],
        ["오탐률 (False Positive)", "15%", "5%", "2%"],
        ["실시간 처리 지연시간", "200ms", "80ms", "30ms"],
        ["은행 간 연합학습", "미도입", "3사 파일럿", "전 금융권 확대"],
    ]
    add_kakao_table(doc, headers, rows)

    add_formatted_paragraph(doc, "※ 상기 수치는 업계 벤치마크 및 카카오뱅크 내부 시뮬레이션 기반 추정치입니다.",
                            font_size=8, color=MID_GRAY, space_before=4)


def build_chapter1(doc):
    """Chapter 1: FDS Technology SOTA Survey."""
    add_page_break(doc)
    add_heading_styled(doc, "Chapter 1. FDS 기술 SOTA 동향", level=1)

    # 1-1. GNN
    add_heading_styled(doc, "1-1. GNN (Graph Neural Network) 기반 관계형 탐지", level=2)
    add_formatted_paragraph(
        doc,
        "대포통장을 활용한 자금 세탁 및 보이스피싱 네트워크는 복잡한 점조직 형태를 띠므로 "
        "행(Row) 기반의 전통적인 머신러닝(RF, XGBoost)으로는 탐지하기 어렵습니다. 이를 극복하기 위해 "
        "노드(계좌/고객)와 엣지(거래/송금) 간의 기하학적 관계를 임베딩하는 GNN이 필수적입니다.",
        font_size=10, space_after=6
    )

    add_bullet_point(doc, "대규모 금융 네트워크에서 이웃 노드의 피처를 샘플링·집계하여 타겟 노드 임베딩을 생성합니다. "
                     "신규 계좌(Unseen nodes)에 대한 귀납적(Inductive) 추론이 가능하여, 실시간 유입 대포통장 탐지에 탁월합니다.",
                     bold_prefix="GraphSAGE:")
    add_bullet_point(doc, "이웃 노드의 정보 집계 시 Attention 가중치를 차등 부여합니다. "
                     "정상 소액 결제(노이즈)와 사기 연루 송금(핵심 단서)을 구분하는 데 GraphSAGE 대비 약 5~8% 높은 AUC를 보입니다.",
                     bold_prefix="GAT (Graph Attention Network):")
    add_bullet_point(doc, "'계좌-송금-계좌' 외에 '기기(IP/MAC)-접속-계좌', '전화번호-연결-고객' 등 "
                     "이기종 그래프 데이터를 한 번에 학습하여 멀티채널 사기를 조기 적발합니다.",
                     bold_prefix="Heterogeneous GNN:")

    # 1-2. Federated Learning
    add_heading_styled(doc, "1-2. Federated Learning (연합학습)", level=2)
    add_formatted_paragraph(
        doc,
        "사기범들은 카카오뱅크에서 토스뱅크, 신한은행을 거쳐 자금을 세탁합니다. "
        "그러나 개인정보보호법으로 인해 은행 간 원본 데이터 공유가 불가능한 '데이터 사일로' 현상이 발생합니다. "
        "연합학습은 각 은행이 로컬에서 자체 FDS 모델을 학습시키고, 원본 데이터가 아닌 "
        "'학습된 가중치(Weight) 업데이트' 정보만을 중앙 서버로 전송하여 글로벌 모델을 만드는 방식입니다. "
        "경쟁 은행의 고객 정보 유출 없이 타행 사기 패턴을 자행 FDS에 즉시 반영(Defense-in-depth)할 수 있습니다.",
        font_size=10, space_after=6
    )

    # 1-3. Real-time Streaming
    add_heading_styled(doc, "1-3. Real-time Streaming & CEP", level=2)
    add_formatted_paragraph(
        doc,
        "Apache Flink 기반 순수 스트리밍 엔진을 사용하여 펌뱅킹 이체를 실시간 스코어링합니다. "
        "Kafka를 통해 유입되는 로그를 Flink CEP로 실시간 윈도우 분석하여, "
        "\"로그인 후 3초 내 한도 상향 및 1초 내 전액 이체\"와 같은 매크로 봇 패턴을 50ms 이내에 차단합니다.",
        font_size=10, space_after=6
    )

    # 1-4. LLM 기반 행동분석
    add_heading_styled(doc, "1-4. LLM 기반 행동분석 및 설명 가능성 (XAI)", level=2)
    add_formatted_paragraph(
        doc,
        "고객의 과거 앱 내 메뉴 이동 경로를 텍스트 시퀀스로 변환한 뒤 LLM을 통해 "
        "'고객 고유의 행동 임베딩'을 추출합니다. 원격 제어 앱 조작 시 발생하는 이질적인 시퀀스를 탐지하며, "
        "모델이 이상거래로 판정한 이유를 자연어로 설명하여 분석가에게 제공합니다.",
        font_size=10, space_after=4
    )
    add_formatted_paragraph(
        doc,
        "예시: \"해당 거래는 평소 패턴과 달리 심야에 VPN IP에서 접속되었으며 "
        "가상화폐 거래소로 전액 송금이 시도되어 위험합니다.\"",
        font_size=9, color=MID_GRAY, italic=False, space_after=8,
        first_line_indent=0.5
    )

    # 1-5. Performance comparison table
    add_heading_styled(doc, "1-5. 기술별 성능 비교표 (공개 벤치마크 기준)", level=2)

    headers = ["데이터셋", "모델(알고리즘)", "정밀도", "재현율", "F1-Score", "AUC-ROC", "Latency"]
    rows = [
        ["Elliptic (Crypto)", "XGBoost", "0.810", "0.720", "0.762", "0.840", "~10ms"],
        ["Elliptic (Crypto)", "GATv2 (SOTA)", "0.885", "0.952", "0.917", "0.920", "~45ms"],
        ["IEEE-CIS (E-com)", "LightGBM", "0.890", "0.850", "0.869", "0.940", "~5ms"],
        ["IEEE-CIS (E-com)", "FinGuard-GNN", "0.921", "0.915", "0.918", "0.981", "~35ms"],
        ["PaySim (Mobile)", "Random Forest", "0.910", "0.820", "0.862", "0.855", "~8ms"],
        ["PaySim (Mobile)", "GraphSAGE+RL", "0.965", "0.981", "0.973", "0.992", "~50ms"],
    ]
    add_kakao_table(doc, headers, rows)

    add_formatted_paragraph(doc, "출처: IEEE, KDD, Nature 등 공개 벤치마크 데이터 기반 성능 비교",
                            font_size=8, color=MID_GRAY, space_before=4)


def build_chapter2(doc):
    """Chapter 2: Competitive Analysis."""
    add_page_break(doc)
    add_heading_styled(doc, "Chapter 2. 경쟁사 분석", level=1)

    add_heading_styled(doc, "2-1. 인터넷전문은행 3사 FDS 비교", level=2)
    add_formatted_paragraph(
        doc,
        "국내 인터넷전문은행 3사(카카오뱅크, 토스뱅크, 케이뱅크)는 각각 차별화된 FDS 전략을 운영하고 있습니다. "
        "카카오뱅크는 자체 내재화와 MLOps 중심, 토스뱅크는 사용자 UX와 외부 연동 중심, "
        "케이뱅크는 KT 통신-금융 결합 탐지에 강점을 보입니다.",
        font_size=10, space_after=8
    )

    headers = ["구분", "카카오뱅크", "토스뱅크", "케이뱅크"]
    rows = [
        ["FDS 철학", "자체 내재화 및 MLOps", "사용자 UX 및 외부 연동", "통신-금융 결합 탐지"],
        ["특화 기술\n(주요)", "AI 셀카 신분증 대조\nSurPASS 명의도용 방지\n딥러닝 통화 패턴 분석",
         "사기의심사이렌 연동\n악성 앱 실시간 탐지/삭제\n기기/통합 신호 스코어링",
         "KT AI 보이스피싱 탐지\n스미싱 문자 AI 판독\n영상통화 2중 검증"],
        ["피해 보상", "선제적 차단 중심", "안심보상제(중고거래 포함)", "명의도용 전액 보상(2025)"],
        ["MLOps 수준", "자동화 파이프라인 운영", "외부 파트너 연동", "KT 인프라 활용"],
        ["강점 요약", "내부 기술 축적·확장성", "빠른 외부 연동·UX", "통신데이터 시너지"],
    ]
    add_kakao_table(doc, headers, rows, col_widths=[2.5, 4.0, 4.0, 4.0])

    add_heading_styled(doc, "2-2. 시중은행 FDS 운영 현황", level=2)

    headers2 = ["은행", "FDS 특징", "핵심 기술"]
    rows2 = [
        ["우리은행", "2026년 AI FDS 검사시스템 고도화", "LLM 기반 RAG 활용 검사 수작업 자동화"],
        ["신한은행", "차세대 '더 넥스트' 기반 AI 하이브리드 FDS", "빅데이터 기반 오탐률 최소화"],
        ["하나은행", "FFT + CNN 결합 보이스피싱 포착", "주파수 분석 기반 실시간 음성 탐지"],
        ["KB국민은행", "내부통제 FDS로 영역 확장", "내부 임직원 부정 거래 감시"],
    ]
    add_kakao_table(doc, headers2, rows2, col_widths=[3.0, 5.5, 6.0])

    add_formatted_paragraph(
        doc,
        "카카오뱅크는 자체 MLOps 파이프라인과 AI 기술 내재화에서 경쟁 우위를 보유하고 있으나, "
        "시중은행들의 LLM 기반 RAG, FFT+CNN 등 신기술 도입이 가속화되고 있어 "
        "지속적인 기술 투자가 필수적입니다.",
        font_size=10, space_before=8, space_after=6
    )


def build_chapter3(doc):
    """Chapter 3: Regulatory Trends."""
    add_page_break(doc)
    add_heading_styled(doc, "Chapter 3. 규제 동향", level=1)

    add_heading_styled(doc, "3-1. 망분리 규제 완화 (2024.08 발표)", level=2)
    add_formatted_paragraph(
        doc,
        "2026년 '디지털 금융보안법' 제정을 통해 물리적 망분리가 원칙적으로 폐지됩니다. "
        "이에 따라 FDS가 실시간 지능형 방어선 역할을 수행하게 되며, "
        "클라우드 기반의 유연한 보안 아키텍처 전환이 가속화됩니다. "
        "금융사는 자체 보안 역량 강화와 함께 실시간 위협 탐지 체계를 필수적으로 구축해야 합니다.",
        font_size=10, space_after=8
    )

    add_heading_styled(doc, "3-2. 전자금융거래법 개정 (2024.09 시행)", level=2)
    add_formatted_paragraph(
        doc,
        "선불업/PG사 규제가 강화되었으며 FDS 고도화 가이드라인 준수가 의무화되었습니다. "
        "사고 발생 시 금융사 책임이 크게 강화되어, 사전 예방적 FDS 투자의 중요성이 더욱 부각됩니다.",
        font_size=10, space_after=8
    )

    add_heading_styled(doc, "3-3. AI 공정성 및 설명 가능성 가이드라인", level=2)
    add_formatted_paragraph(
        doc,
        "AI 알고리즘의 공정성 및 설명 가능성(XAI) 요구가 증대하고 있습니다. "
        "FDS 모델이 특정 연령층이나 집단에 대해 부당한 거래 거절을 하지 않도록 "
        "주기적인 편향성 감사가 필요하며, 차단 사유를 고객에게 투명하게 안내하되 "
        "보안을 위해 세부 로직은 숨기는 균형 잡힌 XAI 전략이 요구됩니다.",
        font_size=10, space_after=8
    )

    # Regulatory roadmap table
    add_heading_styled(doc, "3-4. 규제 로드맵 종합", level=2)
    headers = ["연도", "규제/제도", "핵심 내용", "FDS 영향"]
    rows = [
        ["2024.08", "망분리 규제 완화 발표", "물리적 망분리 원칙적 폐지 예고", "클라우드 FDS 전환 근거"],
        ["2024.09", "전자금융거래법 개정", "선불업/PG 규제 강화, FDS 의무화", "투자 당위성 확보"],
        ["2025", "AI 공정성 가이드라인", "XAI 요구, 편향성 감사", "모델 설명 기능 필수"],
        ["2026", "디지털 금융보안법 제정", "망분리 완전 폐지, 지능형 방어 의무", "실시간 FDS 필수화"],
        ["2027~", "글로벌 AI 규제 조화", "EU AI Act 등 국제 표준 조율", "글로벌 기준 FDS 구축"],
    ]
    add_kakao_table(doc, headers, rows, col_widths=[2.0, 3.5, 4.5, 4.5])

    add_formatted_paragraph(doc, "출처: 금융위원회 '금융분야 망분리 개선 로드맵' (2024.08.13) 등",
                            font_size=8, color=MID_GRAY, space_before=4)


def build_chapter4(doc):
    """Chapter 4: Business Strategy - FDS Investment ROI Analysis."""
    add_page_break(doc)
    add_heading_styled(doc, "Chapter 4. 비즈니스 전략 — FDS 투자 ROI 분석", level=1)

    add_heading_styled(doc, "4-1. 현황 진단", level=2)
    add_formatted_paragraph(
        doc,
        "카카오뱅크는 현재 연간 약 200억 원의 이상거래 관련 피해(직접 손실 + 보상비 + 운영비)가 발생하고 있는 것으로 "
        "추정됩니다. 기존 룰 기반 FDS의 오탐률이 15%에 달해 분석가의 업무 부하가 가중되고 있으며, "
        "신종 사기 유형에 대한 탐지 공백이 확대되고 있습니다.",
        font_size=10, space_after=8
    )

    add_heading_styled(doc, "4-2. 투자 계획 및 ROI 시뮬레이션", level=2)
    add_formatted_paragraph(
        doc,
        "GNN + Federated Learning 기반 차세대 FDS 구축에 3년간 총 150억 원을 투자하여, "
        "연간 사기 피해액을 200억 원에서 50억 원으로 감축(75% 절감)하는 것을 목표로 합니다.",
        font_size=10, space_after=6
    )

    # ROI simulation table
    headers = ["항목", "2026 (1차년)", "2027 (2차년)", "2028 (3차년)", "합계"]
    rows = [
        ["FDS 투자액", "70억 원", "50억 원", "30억 원", "150억 원"],
        ["사기 피해액 (기존)", "200억 원", "200억 원", "200억 원", "600억 원"],
        ["사기 피해액 (예상)", "150억 원", "100억 원", "50억 원", "300억 원"],
        ["연간 절감액", "50억 원", "100억 원", "150억 원", "300억 원"],
        ["순 절감액 (절감-투자)", "-20억 원", "50억 원", "120억 원", "150억 원"],
        ["누적 ROI", "-20억 원", "30억 원", "150억 원", "—"],
    ]
    add_kakao_table(doc, headers, rows)

    add_formatted_paragraph(
        doc,
        "※ 2차년도부터 순 절감액이 양(+)으로 전환되며, 3차년도 누적 기준 150억 원의 순 이익을 달성합니다. "
        "오탐률 감소에 따른 분석가 업무 효율화(인건비 절감 연 10억 원 추정)를 포함하면 "
        "총 ROI는 약 180억 원으로 상향됩니다.",
        font_size=9, color=MID_GRAY, space_before=4, space_after=8
    )

    add_heading_styled(doc, "4-3. 경쟁 차별화 전략", level=2)

    add_bullet_point(doc, "GNN 기반 '금융 지형도'를 업계 최초로 구축하여 대포통장 네트워크를 실시간으로 "
                     "가시화합니다. 이는 기존 룰 기반 시스템 대비 탐지율 30% 이상 향상을 의미합니다.",
                     bold_prefix="[기술 선점]")
    add_bullet_point(doc, "3사(카카오·토스·케이) 연합학습 파일럿을 통해 업계 첫 '프라이버시 보존형 "
                     "은행 간 사기 정보 공유 체계'를 구축합니다.",
                     bold_prefix="[생태계 리더십]")
    add_bullet_point(doc, "XAI 기반 차단 사유 투명 안내로 고객 불만을 50% 감소시키고, "
                     "안심보상제와 결합하여 '가장 안전한 은행' 브랜드 포지셔닝을 강화합니다.",
                     bold_prefix="[고객 신뢰]")

    # Annual savings detail table
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_heading_styled(doc, "4-4. 연간 절감 시뮬레이션 상세", level=2)

    headers2 = ["절감 항목", "현재 비용", "목표 비용", "연간 절감액", "비고"]
    rows2 = [
        ["직접 사기 피해", "120억 원", "30억 원", "90억 원", "GNN+FL 탐지율 개선"],
        ["고객 보상금", "30억 원", "8억 원", "22억 원", "선제 차단 강화"],
        ["분석가 인건비", "25억 원", "15억 원", "10억 원", "오탐률 감소 → 업무 효율화"],
        ["시스템 운영비", "15억 원", "12억 원", "3억 원", "클라우드 최적화"],
        ["규제 과징금 리스크", "10억 원", "2억 원", "8억 원", "컴플라이언스 강화"],
        ["합계", "200억 원", "67억 원", "133억 원", "목표 달성 시 (3차년도)"],
    ]
    add_kakao_table(doc, headers2, rows2)


def build_chapter5(doc):
    """Chapter 5: Future Scenarios 2026-2030."""
    add_page_break(doc)
    add_heading_styled(doc, "Chapter 5. 미래 시나리오 2026-2030", level=1)

    add_formatted_paragraph(
        doc,
        "2026년부터 2030년까지 금융 사기는 AI 기술의 발전과 함께 더욱 정교해질 것으로 예상됩니다. "
        "본 장에서는 3대 핵심 위협 시나리오와 발생확률, 그리고 각 시나리오에 대한 대응 전략을 제시합니다.",
        font_size=10, space_after=8
    )

    # Scenario A
    add_heading_styled(doc, "시나리오 A: 딥페이크 + 합성ID 대량 공격", level=2)
    headers = ["항목", "내용"]
    rows = [
        ["위협수준", "상 (High)"],
        ["발생확률", "85%"],
        ["예상 피해", "연간 수백억 원 규모의 대포통장 양산"],
        ["공격 개요", "AI로 유령 인물(Synthetic ID)의 가짜 신분증 수만 개 생성\n실시간 딥페이크로 영상통화 인증 우회"],
        ["대응 전략", "픽셀 주파수 분석(Anti-Spoofing)\n혈류 감지 Liveness Detection\n터치 압력 기반 행동 생체인증"],
    ]
    add_kakao_table(doc, headers, rows, col_widths=[3.5, 11.0])

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Scenario B
    add_heading_styled(doc, "시나리오 B: AI vs AI 적대적 공방 (Adversarial ML)", level=2)
    headers = ["항목", "내용"]
    rows = [
        ["위협수준", "최상 (Critical)"],
        ["발생확률", "60%"],
        ["예상 피해", "특정 고액 자산가 타겟 APT 공격으로 대규모 자산 유출"],
        ["공격 개요", "FDS 스코어링 역공학(Model Extraction)\n우회 이체 패턴 자동 생성"],
        ["대응 전략", "적대적 훈련(Adversarial Training)\n동적 앙상블 모델 운영\nZero-Trust API 마이크로 세그멘테이션"],
    ]
    add_kakao_table(doc, headers, rows, col_widths=[3.5, 11.0])

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Scenario C
    add_heading_styled(doc, "시나리오 C: CBDC + 디지털 자산 연계 사기", level=2)
    headers = ["항목", "내용"]
    rows = [
        ["위협수준", "중 (Medium)"],
        ["발생확률", "75%"],
        ["예상 피해", "스마트 컨트랙트 허점 이용 예금토큰 탈취 및 해외 세탁"],
        ["공격 개요", "CBDC 시범 사업 및 STO 확산 시 규제 공백 이용"],
        ["대응 전략", "온/오프체인 데이터 퓨전 분석\n스마트 컨트랙트 실시간 해시 감사\n글로벌 가상자산 FDS 연동"],
    ]
    add_kakao_table(doc, headers, rows, col_widths=[3.5, 11.0])

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Response roadmap
    add_heading_styled(doc, "5-4. 시나리오 대응 로드맵", level=2)
    headers_rm = ["기간", "단기 (1년)", "중기 (3년)", "장기 (5년)"]
    rows_rm = [
        ["딥페이크 대응", "Anti-Spoofing 모듈 도입\nLiveness Detection PoC",
         "행동 생체인증 전면 적용\n딥페이크 탐지 AI 고도화",
         "멀티모달 인증 체계 완성\n글로벌 표준 선도"],
        ["AI vs AI 대응", "Adversarial Training 도입\n모델 보안 감사 체계",
         "동적 앙상블 운영\nZero-Trust API 전환",
         "AI 레드팀 상시 운영\n자율 방어 시스템 구축"],
        ["CBDC/DeFi 대응", "온체인 모니터링 PoC\n규제 동향 추적",
         "온/오프체인 퓨전 분석\nSTO FDS 모듈 개발",
         "글로벌 가상자산 FDS 연동\nCBDC 전용 탐지 엔진"],
        ["공통 인프라", "GPU 클러스터 구축\nMLOps 파이프라인 강화",
         "NPU 가속기 도입\n연합학습 3사 파일럿",
         "엣지 컴퓨팅 FDS\n전 금융권 연합학습"],
    ]
    add_kakao_table(doc, headers_rm, rows_rm, col_widths=[2.5, 4.0, 4.0, 4.0])


def build_chapter6(doc):
    """Chapter 6: Implementation Plan."""
    add_page_break(doc)
    add_heading_styled(doc, "Chapter 6. 구현 로드맵", level=1)

    add_heading_styled(doc, "6-1. 단계별 추진 계획", level=2)

    headers = ["Phase", "기간", "주요 과제", "산출물", "예산"]
    rows = [
        ["Phase 1\n(기반 구축)", "2026 Q1-Q2", "GNN 모델 개발\n피처 스토어 구축\nGPU 인프라 확보",
         "GNN PoC 모델\n피처 스토어 v1.0\n인프라 구축 완료", "35억 원"],
        ["Phase 2\n(파일럿)", "2026 Q3-Q4", "GNN+FL 통합 모델\n3사 연합학습 PoC\nXAI 모듈 개발",
         "통합 FDS v1.0\nFL PoC 보고서\nXAI 프로토타입", "35억 원"],
        ["Phase 3\n(확산)", "2027 Q1-Q4", "전면 배포\nMLOps 자동화\n실시간 CEP 연동",
         "FDS v2.0 운영\nMLOps 파이프라인\n대시보드 v2.0", "50억 원"],
        ["Phase 4\n(최적화)", "2028 Q1-Q4", "NPU 가속\n오탐률 2% 달성\n글로벌 확장",
         "FDS v3.0\n성능 최적화 보고서\n글로벌 연동 체계", "30억 원"],
    ]
    add_kakao_table(doc, headers, rows, col_widths=[2.5, 2.5, 3.5, 3.5, 2.0])

    add_heading_styled(doc, "6-2. 소요 인력", level=2)

    headers2 = ["직군", "인원", "역할", "비고"]
    rows2 = [
        ["ML 엔지니어", "8명", "GNN/FL 모델 개발 및 학습", "석·박사급 3명 포함"],
        ["데이터 엔지니어", "5명", "피처 스토어, 데이터 파이프라인", "Kafka/Flink 전문"],
        ["MLOps 엔지니어", "4명", "모델 서빙, 모니터링, CI/CD", "K8s/EKS 전문"],
        ["보안 전문가", "3명", "Adversarial ML, 침투 테스트", "Red Team 운영"],
        ["FDS 분석가", "6명", "룰 관리, 오탐 분석, 고객 대응", "현업 전문가"],
        ["PM/PO", "2명", "프로젝트 관리, 이해관계자 조율", "금융 IT 경험"],
        ["합계", "28명", "—", "—"],
    ]
    add_kakao_table(doc, headers2, rows2)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    add_heading_styled(doc, "6-3. 의사결정 매트릭스", level=2)

    headers3 = ["기술 옵션", "탐지 효과", "구현 난이도", "비용", "시급성", "종합 점수\n(1-10)"]
    rows3 = [
        ["GNN 기반 관계형 탐지", "★★★★★", "★★★★", "★★★", "★★★★★", "9.2"],
        ["Federated Learning", "★★★★", "★★★★★", "★★★★", "★★★★", "8.5"],
        ["Real-time CEP (Flink)", "★★★★", "★★★", "★★★", "★★★★★", "8.8"],
        ["LLM 기반 XAI", "★★★", "★★★★", "★★★★", "★★★", "7.5"],
        ["딥페이크 탐지", "★★★★★", "★★★★★", "★★★★★", "★★★★", "8.0"],
        ["행동 생체인증", "★★★★", "★★★★", "★★★", "★★★", "7.2"],
    ]
    add_kakao_table(doc, headers3, rows3)

    add_formatted_paragraph(doc, "※ ★ 1개 = 낮음(쉬움/저렴), ★ 5개 = 높음(어려움/고가). 종합 점수는 가중 평균.",
                            font_size=8, color=MID_GRAY, space_before=4)


def build_chapter7(doc):
    """Chapter 7: MLOps & Infrastructure."""
    add_page_break(doc)
    add_heading_styled(doc, "Chapter 7. MLOps & 인프라", level=1)

    add_heading_styled(doc, "7-1. MLOps 생명주기 최적화", level=2)
    add_formatted_paragraph(
        doc,
        "FDS 모델은 사기 패턴의 변화(Concept Drift)에 매우 민감합니다. "
        "카카오뱅크는 모델의 수명을 관리하기 위해 자동화된 MLOps 파이프라인을 운영합니다.",
        font_size=10, space_after=6
    )

    add_bullet_point(doc, "내부 거래 로그, 신용정보원 데이터, 통신사 패턴 데이터를 "
                     "실시간 피처 스토어(Feature Store)에 적재합니다.",
                     bold_prefix="[데이터 엔지니어링]")
    add_bullet_point(doc, "매일 발생하는 신규 사기 패턴을 라벨링하여 모델을 매주 또는 매일 업데이트합니다. "
                     "챔피언-챌린저(Champion-Challenger) 테스트를 통한 검증 후 배포합니다.",
                     bold_prefix="[지속적 재학습]")
    add_bullet_point(doc, "AUC 및 오탐률을 실시간 대시보드로 관제합니다. "
                     "성능 저하 감지 시 즉시 롤백 및 알람이 발생합니다.",
                     bold_prefix="[성능 모니터링]")

    # MLOps pipeline table
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    headers = ["단계", "도구/기술", "주기", "담당"]
    rows = [
        ["데이터 수집", "Kafka, Spark Streaming", "실시간", "데이터 엔지니어"],
        ["피처 생성", "Feature Store (Feast)", "실시간/배치", "ML 엔지니어"],
        ["모델 학습", "PyTorch Geometric, DGL", "주간/일간", "ML 엔지니어"],
        ["모델 검증", "Champion-Challenger A/B", "배포 전", "MLOps"],
        ["모델 서빙", "Triton, TorchServe, K8s", "실시간", "MLOps"],
        ["모니터링", "Grafana, Prometheus", "실시간", "MLOps + 분석가"],
        ["롤백", "Blue-Green Deployment", "성능 저하 시", "MLOps"],
    ]
    add_kakao_table(doc, headers, rows)

    add_heading_styled(doc, "7-2. GPU/NPU 기반 가속 인프라", level=2)
    add_formatted_paragraph(
        doc,
        "GNN 및 Transformer 기반 모델은 연산량이 방대합니다. NVIDIA A100/H100 GPU를 활용한 "
        "분산 학습 환경을 구축하고, 추론 단계에서는 Latency 최소화를 위해 전용 NPU 가속기 도입을 고려합니다.",
        font_size=10, space_after=6
    )

    add_formatted_paragraph(
        doc,
        "수천 명의 동시 접속자가 이체 버튼을 누르는 순간, 쿠버네티스(EKS) 환경에서 "
        "모델 서빙 노드가 오토스케일링되어 트래픽을 처리합니다. "
        "목표 Latency는 추론 30ms, 전체 파이프라인 80ms 이내입니다.",
        font_size=10, space_after=8
    )

    # Infrastructure table
    headers2 = ["구성 요소", "현재", "Phase 2 목표", "Phase 4 목표"]
    rows2 = [
        ["학습 인프라", "GPU 4장 (V100)", "GPU 16장 (A100)", "GPU 32장 (H100)"],
        ["추론 인프라", "CPU 기반", "GPU 8장 (A100)", "NPU 전용 클러스터"],
        ["오케스트레이션", "Docker Compose", "EKS (K8s)", "EKS + Service Mesh"],
        ["모니터링", "수동 점검", "Grafana + Prometheus", "AIOps 자동 관제"],
        ["배포 방식", "수동 배포", "Blue-Green", "Canary + Auto-Rollback"],
    ]
    add_kakao_table(doc, headers2, rows2)


def build_chapter8(doc):
    """Chapter 8: Conclusion & Strategic Recommendations."""
    add_page_break(doc)
    add_heading_styled(doc, "Chapter 8. 결론 및 전략적 제언", level=1)

    add_heading_styled(doc, "8-1. 기술적 우위 확보 (Tech Leadership)", level=2)
    add_bullet_point(doc, "단순 계좌 관계를 넘어 가상자산 지갑 주소, 다크웹 유출 정보까지 연결한 "
                     "거대 그래프 데이터베이스를 구축합니다. 이를 통해 은행의 FDS가 단순 룰 엔진에서 "
                     "'금융 범죄 인텔리전스 플랫폼'으로 진화합니다.",
                     bold_prefix="GNN 기반 '금융 지형도' 구축:")
    add_bullet_point(doc, "라벨링되지 않은 대량의 정상 거래 데이터에서 '정상의 정의'를 학습하여, "
                     "정의되지 않은 변칙 사기를 잡아내는 비지도 학습 비중을 확대합니다.",
                     bold_prefix="Self-Supervised Learning:")

    add_heading_styled(doc, "8-2. 비즈니스 신뢰 구축 (Trust & Safety)", level=2)
    add_bullet_point(doc, "고객에게 차단 사유를 투명하게 안내하되, 보안을 위해 세부 로직은 숨기는 "
                     "균형 있는 XAI 전략이 필요합니다.",
                     bold_prefix="설명 가능한 금융 AI (XAI):")
    add_bullet_point(doc, "모델의 편향성을 주기적으로 감사하여 노년층이나 특정 집단에 대한 "
                     "부당한 거래 거절을 방지합니다.",
                     bold_prefix="책임감 있는 AI 거버넌스:")

    add_heading_styled(doc, "8-3. 생태계 협력 (Ecosystem Expansion)", level=2)
    add_bullet_point(doc, "금융위, 경찰청, KISA뿐만 아니라 글로벌 보안 기업과의 "
                     "실시간 위협 인텔리전스 공유 체계를 강화합니다.",
                     bold_prefix="민-관-군 통합 대응:")

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Strategic recommendations summary table
    add_heading_styled(doc, "8-4. 전략적 제언 종합", level=2)
    headers = ["영역", "핵심 제언", "기대 효과", "우선순위"]
    rows = [
        ["기술 혁신", "GNN+FL 하이브리드 FDS 구축", "탐지율 0.97, 오탐 2%", "최우선"],
        ["투자 효율", "3년 150억 투자 → 300억 절감", "ROI 200%, BEP 2차년도", "최우선"],
        ["규제 대응", "XAI + AI 거버넌스 체계", "규제 리스크 최소화", "높음"],
        ["미래 대비", "3대 시나리오 선제 대응", "2030년까지 방어 체계 완성", "높음"],
        ["인재 확보", "ML/보안 전문가 28명 확보", "지속적 기술 경쟁력 유지", "중간"],
        ["생태계", "은행 간 FL, 민-관-군 협력", "업계 리더십 확보", "중간"],
    ]
    add_kakao_table(doc, headers, rows)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Closing statement
    add_formatted_paragraph(
        doc,
        "카카오뱅크는 AI 기반 FDS 고도화를 통해 '가장 안전한 디지털 은행'이라는 비전을 실현할 수 있습니다. "
        "GNN과 Federated Learning의 결합은 단순히 기술 도입을 넘어, "
        "금융 범죄 대응의 패러다임을 전환하는 전략적 투자입니다. "
        "본 보고서의 로드맵을 충실히 이행한다면, 2030년까지 연간 사기 피해를 50억 원 이하로 감축하고 "
        "업계 최고 수준의 FDS 역량을 확보할 수 있을 것입니다.",
        font_size=11, bold=True, color=KAKAO_DARK, space_before=8, space_after=6,
        alignment=WD_ALIGN_PARAGRAPH.LEFT
    )


def build_appendix(doc):
    """Appendix: Glossary + References."""
    add_page_break(doc)
    add_heading_styled(doc, "Appendix. 부록", level=1)

    add_heading_styled(doc, "A-1. 용어 해설 (Glossary)", level=2)

    headers = ["용어", "설명"]
    rows = [
        ["FDS (Fraud Detection System)", "이상거래탐지시스템. 실시간으로 금융 거래를 모니터링하여 사기 거래를 탐지·차단하는 시스템"],
        ["GNN (Graph Neural Network)", "그래프 신경망. 노드와 엣지로 구성된 그래프 데이터를 학습하는 딥러닝 모델"],
        ["GAT (Graph Attention Network)", "그래프 어텐션 네트워크. 이웃 노드에 차등 가중치를 부여하는 GNN 변형"],
        ["GraphSAGE", "대규모 그래프에서 이웃 노드 피처를 샘플링·집계하여 임베딩을 생성하는 GNN 알고리즘"],
        ["Federated Learning (FL)", "연합학습. 데이터를 공유하지 않고 모델 가중치만 교환하여 협력 학습하는 방식"],
        ["CEP (Complex Event Processing)", "복합 이벤트 처리. 실시간 스트리밍 데이터에서 복잡한 패턴을 탐지하는 기술"],
        ["XAI (Explainable AI)", "설명 가능한 인공지능. AI 모델의 판단 근거를 사람이 이해할 수 있게 설명하는 기술"],
        ["MLOps", "ML 모델의 개발-배포-모니터링 전 생명주기를 자동화하는 DevOps 방법론"],
        ["AUC-ROC", "모델 성능 지표. 1에 가까울수록 사기/정상 분류 성능이 우수"],
        ["Concept Drift", "개념 표류. 시간이 지남에 따라 데이터 분포가 변화하는 현상"],
        ["Champion-Challenger", "현행 모델(Champion)과 신규 모델(Challenger)을 A/B 테스트하는 방식"],
        ["Adversarial Training", "적대적 훈련. 모델 공격 샘플을 학습에 포함시켜 견고성을 높이는 방법"],
        ["CBDC (Central Bank Digital Currency)", "중앙은행 디지털 화폐"],
        ["STO (Security Token Offering)", "증권형 토큰 공개. 블록체인 기반 디지털 증권 발행"],
        ["NPU (Neural Processing Unit)", "신경망 처리 전용 프로세서. AI 추론 가속에 특화"],
    ]
    add_kakao_table(doc, headers, rows, col_widths=[5.5, 10.0])

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    add_heading_styled(doc, "A-2. 참고문헌 (References)", level=2)

    refs = [
        "[1] IEEE, \"Federated Graph Neural Networks for Fraud Detection\", 2023.",
        "[2] KDD, \"Inductive Representation Learning on Large Graphs\", Stanford University.",
        "[3] Nature, \"Anti-Money Laundering in Bitcoin: GCN Experiments\".",
        "[4] 금융위원회, \"금융분야 망분리 개선 로드맵\", 2024.08.13.",
        "[5] 금융보안원, \"인공지능 보안 가이드라인\", 2023.",
        "[6] Gartner, \"Top Strategic Technology Trends in Banking 2025\", 2024.",
        "[7] MIT Sloan, \"The Future of MLOps in High-Stakes Environments\".",
        "[8] 카카오뱅크, 토스뱅크, 케이뱅크 2023-2024 경영실적 보고서 및 테크 블로그.",
        "[9] 전자금융거래법 개정안 (2024.09 시행).",
        "[10] EU AI Act, 2024.",
    ]

    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(ref)
        set_run_font(run, size=9, color=MID_GRAY)


def setup_headers_footers(doc):
    """Set up header and footer for the document."""
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = hp.add_run("KakaoBank · AI FDS 고도화 보고서 | CONFIDENTIAL")
        set_run_font(run, size=8, color=MID_GRAY)

        # Footer with page number
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run_pre = fp.add_run("— ")
        set_run_font(run_pre, size=8, color=MID_GRAY)

        # Add page number field
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run_pg = fp.add_run()
        run_pg._r.append(fldChar1)

        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run_pg2 = fp.add_run()
        run_pg2._r.append(instrText)

        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run_pg3 = fp.add_run()
        run_pg3._r.append(fldChar2)

        run_post = fp.add_run(" —")
        set_run_font(run_post, size=8, color=MID_GRAY)


def setup_styles(doc):
    """Configure document styles."""
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(10)
    style.font.color.rgb = DARK_GRAY
    # Set East Asian font
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)

    # Heading styles
    for i in range(1, 5):
        style_name = f'Heading {i}'
        if style_name in doc.styles:
            h_style = doc.styles[style_name]
            h_style.font.name = FONT_NAME
            h_style.font.color.rgb = KAKAO_DARK
            rPr = h_style.element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:eastAsia'), FONT_NAME)

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    doc = Document()
    setup_styles(doc)

    # Build all sections
    build_cover_page(doc)
    build_toc(doc)
    build_executive_summary(doc)
    build_chapter1(doc)
    build_chapter2(doc)
    build_chapter3(doc)
    build_chapter4(doc)
    build_chapter5(doc)
    build_chapter6(doc)
    build_chapter7(doc)
    build_chapter8(doc)
    build_appendix(doc)

    # Setup headers and footers
    setup_headers_footers(doc)

    # Save
    doc.save(OUTPUT_FILE)
    print(f"[OK] Report saved: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
