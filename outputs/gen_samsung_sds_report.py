"""
Samsung SDS COBOL → Cloud Migration Report Generator
Client-delivery quality DOCX with proper headers, footers, tables, and TOC.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree
import copy

# ─── Constants ───
SAMSUNG_BLUE = RGBColor(0x14, 0x28, 0xA0)
SAMSUNG_BLUE_HEX = "1428A0"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)
LIGHT_GRAY_HEX = "F2F2F2"
FONT_NAME = "맑은 고딕"
FONT_NAME_EN = "Malgun Gothic"

OUTPUT_DIR = r"C:\dev\cmux-win\outputs"
OUTPUT_DOCX = os.path.join(OUTPUT_DIR, "SamsungSDS_COBOL_Migration_Report.docx")
OUTPUT_PDF = os.path.join(OUTPUT_DIR, "SamsungSDS_COBOL_Migration_Report.pdf")


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set individual cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            border = parse_xml(
                f'<w:{edge} {nsdecls("w")} w:val="{val["val"]}" w:sz="{val["sz"]}" '
                f'w:space="0" w:color="{val["color"]}"/>'
            )
            tcBorders.append(border)
    tcPr.append(tcBorders)


def set_table_borders(table, color="1428A0", sz="4"):
    """Set borders for entire table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def set_font(run, name=FONT_NAME, size=None, bold=False, color=None, italic=False):
    """Apply font settings to a run."""
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_paragraph(doc, text, style=None, alignment=None, font_size=None, bold=False,
                  color=None, space_before=None, space_after=None, font_name=FONT_NAME):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph(style=style)
    if alignment is not None:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, name=font_name, size=font_size, bold=bold, color=color)
    return p


def add_heading_styled(doc, text, level=1):
    """Add a heading with Malgun Gothic font."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        if level == 1:
            run.font.color.rgb = SAMSUNG_BLUE
            run.font.size = Pt(20)
        elif level == 2:
            run.font.color.rgb = SAMSUNG_BLUE
            run.font.size = Pt(16)
        elif level == 3:
            run.font.color.rgb = DARK_GRAY
            run.font.size = Pt(13)
    return h


def add_table_with_style(doc, headers, rows, col_widths=None):
    """Add a professional table with Samsung blue header row and borders."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Set borders
    set_table_borders(table, color=SAMSUNG_BLUE_HEX, sz="6")

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = ""
        set_cell_shading(hdr_cells[i], SAMSUNG_BLUE_HEX)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_font(run, size=10, bold=True, color=WHITE)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = ""
            p = row_cells[c_idx].paragraphs[0]
            run = p.add_run(str(cell_text))
            set_font(run, size=9, color=DARK_GRAY)
            row_cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # Alternate row shading
            if r_idx % 2 == 1:
                set_cell_shading(row_cells[c_idx], LIGHT_GRAY_HEX)

    # Set column widths if specified
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(width)

    return table


def add_bullet_list(doc, items, indent_level=0, font_size=10):
    """Add bullet points."""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        if indent_level > 0:
            p.paragraph_format.left_indent = Cm(indent_level * 1.27)
        run = p.add_run(item)
        set_font(run, size=font_size, color=DARK_GRAY)


def add_numbered_list(doc, items, font_size=10):
    """Add numbered items as regular paragraphs with numbers."""
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(f"{i}. {item}")
        set_font(run, size=font_size, color=DARK_GRAY)


def add_code_block(doc, code_text, language=""):
    """Add a code block with gray background."""
    if language:
        add_paragraph(doc, f"[{language}]", font_size=8, color=RGBColor(0x66, 0x66, 0x66),
                      space_before=6, space_after=0)
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        run.font.color.rgb = DARK_GRAY
        # Set shading for the paragraph
        pPr = p._element.get_or_add_pPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
        pPr.append(shading)


def add_info_box(doc, text, border_color=SAMSUNG_BLUE_HEX):
    """Add an information callout box."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    # Add left border via paragraph borders
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="8" w:color="{border_color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    # Background shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EEF0F8" w:val="clear"/>')
    pPr.append(shading)
    run = p.add_run(text)
    set_font(run, size=9, color=DARK_GRAY, italic=True)


def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()


def setup_header_footer(doc):
    """Set up header and footer for all sections."""
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.clear()

        # Header text
        run = hp.add_run("Samsung SDS")
        set_font(run, size=8, bold=True, color=SAMSUNG_BLUE)

        run = hp.add_run(" · 메인프레임 클라우드 마이그레이션 보고서")
        set_font(run, size=8, color=DARK_GRAY)

        # Add tab and right-aligned CONFIDENTIAL
        run = hp.add_run("\t\t")
        set_font(run, size=8, color=DARK_GRAY)
        run = hp.add_run("CONFIDENTIAL")
        set_font(run, size=8, bold=True, color=RGBColor(0xCC, 0x00, 0x00))

        # Add bottom border to header
        pPr = hp._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="8" w:space="4" w:color="{SAMSUNG_BLUE_HEX}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

        # Tab stops for header (center and right)
        tabs = parse_xml(
            f'<w:tabs {nsdecls("w")}>'
            f'  <w:tab w:val="center" w:pos="4536"/>'
            f'  <w:tab w:val="right" w:pos="9072"/>'
            f'</w:tabs>'
        )
        pPr.append(tabs)

        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.clear()

        # Add top border to footer
        fpPr = fp._element.get_or_add_pPr()
        fBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="4" w:color="{SAMSUNG_BLUE_HEX}"/>'
            f'</w:pBdr>'
        )
        fpPr.append(fBdr)

        # Footer text with page number
        run = fp.add_run("삼성SDS 메인프레임 클라우드 마이그레이션 최종 보고서  |  ")
        set_font(run, size=7, color=RGBColor(0x99, 0x99, 0x99))

        # Page number field
        fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        r1 = parse_xml(f'<w:r {nsdecls("w")}></w:r>')
        r1.append(fld_char_begin)
        fp._element.append(r1)

        instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        r2 = parse_xml(f'<w:r {nsdecls("w")}></w:r>')
        rPr2 = parse_xml(f'<w:rPr {nsdecls("w")}><w:rFonts w:ascii="{FONT_NAME}" w:eastAsia="{FONT_NAME}"/><w:sz w:val="14"/><w:color w:val="999999"/></w:rPr>')
        r2.append(rPr2)
        r2.append(instr)
        fp._element.append(r2)

        fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        r3 = parse_xml(f'<w:r {nsdecls("w")}></w:r>')
        r3.append(fld_char_end)
        fp._element.append(r3)

        run2 = fp.add_run(" / ")
        set_font(run2, size=7, color=RGBColor(0x99, 0x99, 0x99))

        # Total pages field
        fld_char_begin2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        r4 = parse_xml(f'<w:r {nsdecls("w")}></w:r>')
        r4.append(fld_char_begin2)
        fp._element.append(r4)

        instr2 = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>')
        r5 = parse_xml(f'<w:r {nsdecls("w")}></w:r>')
        rPr5 = parse_xml(f'<w:rPr {nsdecls("w")}><w:rFonts w:ascii="{FONT_NAME}" w:eastAsia="{FONT_NAME}"/><w:sz w:val="14"/><w:color w:val="999999"/></w:rPr>')
        r5.append(rPr5)
        r5.append(instr2)
        fp._element.append(r5)

        fld_char_end2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        r6 = parse_xml(f'<w:r {nsdecls("w")}></w:r>')
        r6.append(fld_char_end2)
        fp._element.append(r6)


def add_toc(doc):
    """Add Table of Contents field."""
    add_heading_styled(doc, "목 차", level=1)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)

    # TOC field
    r = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    r._element.append(fldChar1)

    r2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    r2._element.append(instrText)

    r3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    r3._element.append(fldChar2)

    # Placeholder text
    r4 = p.add_run("(목차를 업데이트하려면 이 영역을 우클릭 → '필드 업데이트'를 선택하세요)")
    set_font(r4, size=10, color=RGBColor(0x99, 0x99, 0x99), italic=True)

    r5 = p.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    r5._element.append(fldChar3)


def create_cover_page(doc):
    """Create the Samsung SDS branded cover page."""
    # Remove header/footer from first section (cover page)
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # Clear first page header
    first_header = section.first_page_header
    for p in first_header.paragraphs:
        p.clear()

    # Clear first page footer
    first_footer = section.first_page_footer
    for p in first_footer.paragraphs:
        p.clear()

    # Add spacer
    for _ in range(4):
        add_paragraph(doc, "", font_size=12, space_after=0, space_before=0)

    # Samsung blue top bar (using a 1-cell table)
    bar_table = doc.add_table(rows=1, cols=1)
    bar_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = bar_table.rows[0].cells[0]
    set_cell_shading(cell, SAMSUNG_BLUE_HEX)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(" ")
    run.font.size = Pt(4)
    # Set table width to full page
    tbl = bar_table._tbl
    tblPr = tbl.tblPr
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tblW)

    add_paragraph(doc, "", font_size=24, space_after=0)

    # Title
    add_paragraph(doc, "삼성SDS", alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  font_size=36, bold=True, color=SAMSUNG_BLUE, space_after=0)
    add_paragraph(doc, "메인프레임 클라우드 마이그레이션", alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  font_size=28, bold=True, color=SAMSUNG_BLUE, space_before=6, space_after=0)
    add_paragraph(doc, "최종 보고서", alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  font_size=28, bold=True, color=SAMSUNG_BLUE, space_before=6, space_after=12)

    # Subtitle
    add_paragraph(doc, "AS-IS 분석  ·  TO-BE 설계  ·  실행 계획",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  font_size=14, color=DARK_GRAY, space_after=6)

    # Divider line
    div_table = doc.add_table(rows=1, cols=1)
    div_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    div_cell = div_table.rows[0].cells[0]
    set_cell_shading(div_cell, SAMSUNG_BLUE_HEX)
    div_cell.text = ""
    dp = div_cell.paragraphs[0]
    dr = dp.add_run(" ")
    dr.font.size = Pt(1)
    dtbl = div_table._tbl
    dtblPr = dtbl.tblPr
    dtblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="3000" w:type="pct"/>')
    dtblPr.append(dtblW)

    add_paragraph(doc, "", font_size=18, space_after=0)

    # Company info
    add_paragraph(doc, "삼성SDS 주식회사", alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  font_size=14, bold=True, color=DARK_GRAY, space_after=4)
    add_paragraph(doc, "AgenticAI Corp  ·  AI 융합학과", alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  font_size=11, color=RGBColor(0x66, 0x66, 0x66), space_after=4)
    add_paragraph(doc, "2026년 4월 29일", alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  font_size=11, color=RGBColor(0x66, 0x66, 0x66), space_after=18)

    # Confidential notice
    conf_table = doc.add_table(rows=1, cols=1)
    conf_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    conf_cell = conf_table.rows[0].cells[0]
    set_cell_shading(conf_cell, "FFF3F3")
    conf_cell.text = ""
    cp = conf_cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run("CONFIDENTIAL")
    set_font(cr, size=11, bold=True, color=RGBColor(0xCC, 0x00, 0x00))
    cr2 = cp.add_run("\n본 문서는 삼성SDS 주식회사의 기밀 정보를 포함하고 있습니다.\n무단 복제, 배포, 공개를 금합니다.")
    set_font(cr2, size=8, color=RGBColor(0x99, 0x33, 0x33))
    # Set border for confidential box
    set_cell_borders(conf_cell,
                     top={"val": "single", "sz": "4", "color": "CC0000"},
                     bottom={"val": "single", "sz": "4", "color": "CC0000"},
                     left={"val": "single", "sz": "4", "color": "CC0000"},
                     right={"val": "single", "sz": "4", "color": "CC0000"})
    ctbl = conf_table._tbl
    ctblPr = ctbl.tblPr
    ctblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="3500" w:type="pct"/>')
    ctblPr.append(ctblW)

    # Bottom blue bar
    add_paragraph(doc, "", font_size=6)
    bar2_table = doc.add_table(rows=1, cols=1)
    bar2_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    bar2_cell = bar2_table.rows[0].cells[0]
    set_cell_shading(bar2_cell, SAMSUNG_BLUE_HEX)
    bar2_cell.text = ""
    bp = bar2_cell.paragraphs[0]
    br = bp.add_run(" ")
    br.font.size = Pt(4)
    b2tbl = bar2_table._tbl
    b2tblPr = b2tbl.tblPr
    b2tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    b2tblPr.append(b2tblW)


def set_default_styles(doc):
    """Configure default document styles."""
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(10)
    style.font.color.rgb = DARK_GRAY
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Heading styles
    for level in range(1, 4):
        style_name = f'Heading {level}'
        if style_name in doc.styles:
            hs = doc.styles[style_name]
            hs.font.name = FONT_NAME
            hs.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
            if level == 1:
                hs.font.size = Pt(20)
                hs.font.color.rgb = SAMSUNG_BLUE
                hs.font.bold = True
                hs.paragraph_format.space_before = Pt(24)
                hs.paragraph_format.space_after = Pt(12)
            elif level == 2:
                hs.font.size = Pt(16)
                hs.font.color.rgb = SAMSUNG_BLUE
                hs.font.bold = True
                hs.paragraph_format.space_before = Pt(18)
                hs.paragraph_format.space_after = Pt(8)
            elif level == 3:
                hs.font.size = Pt(13)
                hs.font.color.rgb = DARK_GRAY
                hs.font.bold = True
                hs.paragraph_format.space_before = Pt(12)
                hs.paragraph_format.space_after = Pt(6)

    # List Bullet style
    if 'List Bullet' in doc.styles:
        lb = doc.styles['List Bullet']
        lb.font.name = FONT_NAME
        lb.font.size = Pt(10)
        lb.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)


def build_report():
    """Build the complete Samsung SDS COBOL Migration Report."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    set_default_styles(doc)

    # ═══════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════
    create_cover_page(doc)
    add_page_break(doc)

    # Set up headers and footers
    setup_header_footer(doc)

    # ═══════════════════════════════════════════
    # TABLE OF CONTENTS
    # ═══════════════════════════════════════════
    add_toc(doc)
    add_page_break(doc)

    # ═══════════════════════════════════════════
    # EXECUTIVE SUMMARY
    # ═══════════════════════════════════════════
    add_heading_styled(doc, "보고서 개요", level=1)

    add_paragraph(doc, "본 통합 보고서는 삼성SDS COBOL 메인프레임의 클라우드 전환을 위한 세 단계 산출물을 하나의 의사결정 문서로 통합한 것입니다.",
                  font_size=10, space_after=8)

    add_bullet_list(doc, [
        "1부. AS-IS 시스템 분석 — 삼성SDS COBOL 메인프레임의 모듈 의존성, 데이터 흐름(VSAM/DB2/IMS-DB), 배치 스케줄, CICS 온라인 트랜잭션, 그리고 마이그레이션 위험도 TOP 5.",
        "2부. TO-BE 마이그레이션 코드 패키지 — Strangler Fig 기반 Spring Boot 스켈레톤, DB2→PostgreSQL Flyway 스크립트, Kafka 이벤트 브릿지, Terraform(EKS/RDS/MSK) IaC, GitHub Actions CI/CD, Helm(HPA/PDB/NetworkPolicy) 차트.",
        "3부. 최종 마이그레이션 계획서 — Executive Summary, AS-IS vs TO-BE 비교, 18개월 3단계 로드맵, 5×5 리스크 매트릭스, 5년 TCO 비교, 인력 전환 계획, Go/No-Go 체크리스트."
    ])

    add_info_box(doc, "본 리포트는 삼성SDS의 금융 차세대 표준(Anyframe / F-ERP) 및 Tier-1 금융사(은행·보험) 메인프레임 운영 패턴을 기반으로 작성되었습니다. 삼성SDS의 운영 표준과 한국 금융권 레거시 아키텍처 패턴을 결합한 전문가 수준의 참고 아키텍처입니다.")

    add_page_break(doc)

    # ═══════════════════════════════════════════
    # PART 1: AS-IS SYSTEM ANALYSIS
    # ═══════════════════════════════════════════
    add_heading_styled(doc, "1부 — AS-IS 시스템 분석", level=1)

    # --- 1.1 Module Dependency ---
    add_heading_styled(doc, "1. 모듈 의존성 그래프 (JCL → Program → COPY 멤버)", level=2)

    add_paragraph(doc, "삼성SDS 시스템은 비즈니스 프로세스의 논리적 분리를 위해 엄격한 Prefix 표준을 준수합니다.",
                  font_size=10, space_after=6)

    add_heading_styled(doc, "의존성 계층 구조", level=3)

    add_table_with_style(doc,
        headers=["계층", "약어", "역할"],
        rows=[
            ["JCL (Job Control Language)", "JB / ST", "작업 단위(Job)와 실행 순서(Step) 정의"],
            ["Main Program (COBOL)", "BP", "PROCEDURE DIVISION 내 비즈니스 로직 및 SQL/DLI 처리"],
            ["Copybook (CPY)", "CP", "DB2 DCLGEN 및 공통 통신 영역(COMMAREA) 레이아웃"],
            ["Subroutine", "SUB / SP", "재사용 가능한 공통 서브루틴 모듈"]
        ])

    add_paragraph(doc, "", font_size=4)

    add_heading_styled(doc, "실행 흐름 예시", level=3)
    add_code_block(doc, """JB: 일마감_메인
  ├── ST01: 원장추출 ──► BP: ACC_EXTRACT
  └── ST02: 이자계산 ──► BP: ACC_INTEREST
                            ├── CP: ACC_MASTER_REC
                            ├── CP: INTEREST_RATE_TBL
                            └── SUB: TAX_CALC_CORE""")

    # --- 1.2 Data Flow ---
    add_heading_styled(doc, "2. 데이터 흐름도 (VSAM, DB2, IMS-DB)", level=2)

    add_paragraph(doc, "한국 금융 시스템은 트랜잭션 특성에 따라 세 가지 주요 저장소를 혼용합니다.",
                  font_size=10, space_after=6)

    add_table_with_style(doc,
        headers=["데이터 구분", "저장소 타입", "주요 용도", "흐름 방향"],
        rows=[
            ["고객/상품 원장", "DB2", "관계형 데이터, 복합 쿼리 조회", "Online/Batch ↔ DB2"],
            ["계좌 계층 구조", "IMS-DB", "고속 계층형 트리 조회(조직/계좌)", "Online (Read/Update)"],
            ["거래 로그 / 마감", "VSAM", "대량 순차 입출력 (KSDS, ESDS)", "Online (Write) → Batch (Read)"],
            ["대외 기관 전문", "Anylink", "FEP 연동 및 데이터 동기화", "PGM ↔ External System"]
        ])

    # --- 1.3 Batch Schedule ---
    add_heading_styled(doc, "3. 배치 스케줄 의존성 (EOD / EOM 기준)", level=2)

    add_paragraph(doc, "삼성SDS는 A-Auto 또는 Control-M을 사용해 작업 간 선후행 관계를 관리합니다.",
                  font_size=10, space_after=6)

    add_table_with_style(doc,
        headers=["Phase", "시간", "주요 작업"],
        rows=[
            ["Phase 1 — Closing Prep", "18:00~", "온라인 원장 스냅샷 생성 및 미처리 전문 정리"],
            ["Phase 2 — Main Processing", "22:00~", "Critical Path: 잔액 업데이트 → 이자 계상 → 세금 계산 → 원장 반영"],
            ["Phase 3 — Data Delivery", "02:00~", "DW(Data Warehouse) 전송 및 대외 정산(금결원 등)"]
        ])

    # --- 1.4 CICS Transactions ---
    add_heading_styled(doc, "4. 온라인 트랜잭션 목록 (CICS MAP 기준)", level=2)

    add_paragraph(doc, "CICS BMS(Basic Mapping Support) 화면 제어와 메인 프로그램을 매핑한 핵심 트랜잭션입니다.",
                  font_size=10, space_after=6)

    add_table_with_style(doc,
        headers=["트랜잭션 ID", "서비스명", "BMS MAP ID", "호출 프로그램", "주요 기능"],
        rows=[
            ["BAL1", "잔액 조회", "MAPB100", "OPAA100", "실시간 계좌 잔액 및 거래 가능액 확인"],
            ["TRF1", "계좌 이체", "MAPT200", "OPTA200", "당·타행 실시간 이체 및 전문 송수신"],
            ["LON1", "대출 승인", "MAPL300", "OPLA300", "신용도 조회 및 대출 한도 산출 로직"],
            ["CST1", "고객 정보", "MAPC400", "OPCA400", "IMS-DB 연동 고객 통합 정보 관리"]
        ])

    # --- 1.5 Risk TOP 5 ---
    add_heading_styled(doc, "5. 마이그레이션 위험도 TOP 5 모듈", level=2)

    add_paragraph(doc, "Cloud-Native(Java / Spring) 환경 전환 시 기술적 난이도가 높은 모듈입니다.",
                  font_size=10, space_after=6)

    add_table_with_style(doc,
        headers=["순위", "모듈명", "핵심 위험"],
        rows=[
            ["1", "이자/세금 연산 엔진\n(Interest / Tax Core)", "COBOL의 COMP-3 수치 연산과 Java BigDecimal 간 부동소수점 오차 발생 위험"],
            ["2", "대외 전문 파싱 모듈\n(FEP Handler)", "EBCDIC 코드셋 변환 및 고정 길이 전문의 가변 길이 처리 시 데이터 손실"],
            ["3", "대량 배치 I/O\n(Massive Batch Processing)", "메인프레임 IMS-DB / VSAM 초고속 I/O 성능을 클라우드 DB가 수용 불가 시 배치 지연"],
            ["4", "다중 데이터 소스 트랜잭션\n(2PC Module)", "IMS-DB와 DB2 동시 업데이트 프로그램의 분산 트랜잭션 일관성 보장"],
            ["5", "레거시 리포트 레이아웃\n(Host Report)", "특정 프린터 제어 문자와 폼 형식이 Web/PDF 환경에서 레이아웃 깨짐"]
        ])

    add_page_break(doc)

    # ═══════════════════════════════════════════
    # PART 2: TO-BE MIGRATION CODE PACKAGE
    # ═══════════════════════════════════════════
    add_heading_styled(doc, "2부 — TO-BE 마이그레이션 코드 패키지", level=1)

    add_info_box(doc, "한국 금융권(은행/보험) 메인프레임의 COBOL 자산을 AWS EKS 기반 Spring Boot 마이크로서비스로 점진 전환하기 위한 표준 산출물 묶음입니다.")

    # --- 2.1 Strangler Fig ---
    add_heading_styled(doc, "1. Strangler Fig 패턴 — Spring Boot 스켈레톤", level=2)

    add_paragraph(doc, "레거시 시스템을 점진적으로 대체하기 위해 API Gateway/Facade 레이어에서 트래픽을 분기하는 Strangler Fig 패턴을 적용합니다. 신규 트랜잭션은 Java 서비스, 미전환 영역은 LegacyCobolClient를 통해 호스트로 위임합니다.",
                  font_size=10, space_after=6)

    add_heading_styled(doc, "1-1. 프로젝트 구조", level=3)
    add_code_block(doc, """account-service/
├── src/main/java/com/samsung/sds/finance/
│   ├── AccountApplication.java
│   ├── config/
│   │   └── StranglerRoutingConfig.java        # 라우팅 룰 정의 (Feature Flag)
│   ├── controller/
│   │   └── AccountFacadeController.java       # Strangler Facade API
│   ├── service/
│   │   ├── NewAccountService.java             # 신규 Java 비즈니스 로직
│   │   └── LegacyCobolClient.java             # 레거시 API/FEP 호출
│   ├── domain/
│   │   ├── Account.java
│   │   └── AccountDto.java
│   └── infra/
│       ├── KafkaProducerConfig.java
│       └── ObservabilityConfig.java
├── src/main/resources/
│   ├── application.yml
│   └── db/migration/                          # Flyway 스크립트
└── pom.xml""")

    add_heading_styled(doc, "1-2. AccountFacadeController.java (Strangler Facade)", level=3)
    add_code_block(doc, """@RestController
@RequestMapping("/api/v1/accounts")
public class AccountFacadeController {

    private final NewAccountService newService;
    private final LegacyCobolClient legacyClient;
    private final FeatureFlagService featureFlag;

    @GetMapping("/{accountId}")
    public ResponseEntity<AccountDto> getAccount(@PathVariable String accountId) {
        if (featureFlag.isEnabled("ACCOUNT_READ_NEW", accountId)) {
            return ResponseEntity.ok(newService.findById(accountId));
        }
        return ResponseEntity.ok(legacyClient.fetchAccount(accountId));
    }

    @PostMapping("/{accountId}/transfer")
    public ResponseEntity<TransferResultDto> transfer(
            @PathVariable String accountId, @RequestBody TransferRequest req) {
        if (featureFlag.isEnabled("TRANSFER_NEW", accountId)) {
            return ResponseEntity.ok(newService.transfer(accountId, req));
        }
        return ResponseEntity.ok(legacyClient.callLegacyTransfer(accountId, req));
    }
}""", "Java")

    add_heading_styled(doc, "1-3. StranglerRoutingConfig.java (라우팅 룰)", level=3)
    add_paragraph(doc, "점진 전환 단계: shadow → canary 5% → 25% → 50% → 100%",
                  font_size=10, space_after=4)
    add_code_block(doc, """@Configuration
public class StranglerRoutingConfig {
    @Bean
    public FeatureFlagService featureFlagService(
            @Value("${unleash.url}") String url,
            @Value("${unleash.app-name}") String app) {
        return new UnleashFeatureFlagService(url, app);
    }
}""", "Java")

    # --- 2.2 DB2 → PostgreSQL ---
    add_heading_styled(doc, "2. DB2 → PostgreSQL 전환 — Flyway 마이그레이션 스크립트", level=2)

    add_heading_styled(doc, "2-1. V1__init_account_schema.sql", level=3)
    add_code_block(doc, """-- DB2의 ACCOUNT_MASTER 테이블을 PostgreSQL로 이관
CREATE TABLE account_master (
    account_id      VARCHAR(20)    PRIMARY KEY,
    customer_id     VARCHAR(20)    NOT NULL,
    product_code    VARCHAR(10)    NOT NULL,
    balance         NUMERIC(19, 2) NOT NULL DEFAULT 0,  -- COBOL COMP-3 → NUMERIC
    open_date       DATE           NOT NULL,
    status_cd       CHAR(2)        NOT NULL,
    last_txn_at     TIMESTAMPTZ    NOT NULL DEFAULT now(),
    version         BIGINT         NOT NULL DEFAULT 0   -- 낙관적 락
);
CREATE INDEX idx_account_customer ON account_master (customer_id);
CREATE INDEX idx_account_status   ON account_master (status_cd);""", "SQL")

    add_heading_styled(doc, "2-2. V2__transaction_log_partitioned.sql", level=3)
    add_code_block(doc, """-- VSAM 거래로그 → PostgreSQL 파티션 테이블
CREATE TABLE transaction_log (
    txn_id        BIGINT          NOT NULL,
    account_id    VARCHAR(20)     NOT NULL,
    txn_type      CHAR(2)         NOT NULL,
    amount        NUMERIC(19, 2)  NOT NULL,
    txn_at        TIMESTAMPTZ     NOT NULL,
    PRIMARY KEY (txn_id, txn_at)
) PARTITION BY RANGE (txn_at);

CREATE TABLE transaction_log_2026_q2 PARTITION OF transaction_log
    FOR VALUES FROM ('2026-04-01') TO ('2026-07-01');""", "SQL")

    add_heading_styled(doc, "2-3. V3__cobol_data_type_mapping.sql", level=3)
    add_table_with_style(doc,
        headers=["COBOL 타입", "DB2 타입", "PostgreSQL 타입", "비고"],
        rows=[
            ["PIC S9(15)V99 COMP-3", "DECIMAL(17,2)", "NUMERIC(19,2)", "잔액·금액 필드"],
            ["PIC X(02)", "CHAR(2)", "CHAR(2)", "상태코드 등 고정길이"],
            ["PIC X(20)", "VARCHAR(20)", "VARCHAR(20)", "계좌번호·고객ID"],
            ["PIC 9(8)", "DATE", "DATE", "날짜 필드"]
        ])

    add_heading_styled(doc, "2-4. application.yml (Flyway 설정)", level=3)
    add_code_block(doc, """spring:
  datasource:
    url: jdbc:postgresql://${DB_HOST}:5432/finance
    username: ${DB_USER}
    password: ${DB_PASS}
  flyway:
    enabled: true
    locations: classpath:db/migration
    baseline-on-migrate: true
    out-of-order: false
    validate-on-migrate: true""", "YAML")

    # --- 2.3 Kafka Bridge ---
    add_heading_styled(doc, "3. Kafka 이벤트 브릿지 — COBOL ↔ Java 공존기", level=2)

    add_paragraph(doc, "레거시와 신규 시스템 사이에서 상태 변경 이벤트를 양방향으로 전달하는 CDC + Outbox 패턴 브릿지입니다.",
                  font_size=10, space_after=6)

    add_heading_styled(doc, "아키텍처 개요", level=3)
    add_code_block(doc, """[메인프레임 DB2/VSAM] ──CDC(IBM IIDR/Debezium)──► [Kafka MSK] ──► [Java consumer]
                                                       ▲
[Java account-service] ──outbox 테이블──► [Debezium]──┘
[Java account-service] ──producer──► [Kafka MSK] ──► [MQ Bridge → 호스트 IMS-DB]""")

    add_heading_styled(doc, "AccountEventProducer.java", level=3)
    add_code_block(doc, """@Service
public class AccountEventProducer {
    private final KafkaTemplate<String, AccountEvent> kafka;

    @Transactional
    public void publishBalanceChanged(Account a) {
        AccountEvent ev = AccountEvent.builder()
            .eventId(UUID.randomUUID().toString())
            .accountId(a.getId())
            .balance(a.getBalance())
            .occurredAt(Instant.now())
            .build();
        kafka.send("account.balance-changed.v1", a.getId(), ev);
    }
}""", "Java")

    add_heading_styled(doc, "LegacyEventConsumer.java (CDC 수신)", level=3)
    add_code_block(doc, """@Component
public class LegacyEventConsumer {
    private final NewAccountService newService;

    @KafkaListener(topics = "legacy.cobol.account.cdc",
                   groupId = "account-service", concurrency = "6")
    public void onLegacyChange(@Payload LegacyCdcEvent ev, Acknowledgment ack) {
        try {
            newService.applyLegacyChange(ev);  // EBCDIC→UTF-8, COMP-3→BigDecimal
            ack.acknowledge();
        } catch (DataConversionException e) {
            throw e;  // DLQ 처리 — SeekToCurrentErrorHandler + DLT
        }
    }
}""", "Java")

    # --- 2.4 Terraform ---
    add_heading_styled(doc, "4. Terraform — EKS / RDS / MSK 인프라 코드", level=2)

    add_heading_styled(doc, "4-1. main.tf", level=3)
    add_code_block(doc, """terraform {
  required_version = ">= 1.7.0"
  required_providers {
    aws = { source = "hashicorp/aws", version = "~> 5.50" }
  }
  backend "s3" {
    bucket         = "samsung-sds-tfstate"
    key            = "finance/migration/terraform.tfstate"
    region         = "ap-northeast-2"
    dynamodb_table = "tf-lock"
    encrypt        = true
  }
}""", "HCL")

    add_heading_styled(doc, "4-2. vpc.tf", level=3)
    add_code_block(doc, """module "vpc" {
  source  = "terraform-aws-modules/vpc/aws"
  version = "5.8.1"
  name = "finance-mig-vpc"
  cidr = "10.40.0.0/16"
  azs             = ["ap-northeast-2a", "ap-northeast-2b", "ap-northeast-2c"]
  private_subnets = ["10.40.1.0/24", "10.40.2.0/24", "10.40.3.0/24"]
  public_subnets  = ["10.40.101.0/24", "10.40.102.0/24", "10.40.103.0/24"]
  enable_nat_gateway   = true
  single_nat_gateway   = false
  enable_dns_hostnames = true
}""", "HCL")

    add_heading_styled(doc, "4-3. eks.tf", level=3)
    add_code_block(doc, """module "eks" {
  source          = "terraform-aws-modules/eks/aws"
  version         = "20.13.1"
  cluster_name    = "finance-mig"
  cluster_version = "1.30"
  vpc_id     = module.vpc.vpc_id
  subnet_ids = module.vpc.private_subnets
  cluster_endpoint_private_access = true
  cluster_endpoint_public_access  = false
  eks_managed_node_groups = {
    app = {
      desired_size   = 3
      min_size       = 3
      max_size       = 12
      instance_types = ["m6i.xlarge"]
      capacity_type  = "ON_DEMAND"
    }
  }
  enable_irsa = true
}""", "HCL")

    add_heading_styled(doc, "4-4. rds.tf (Aurora PostgreSQL)", level=3)
    add_code_block(doc, """module "rds_postgres" {
  source         = "terraform-aws-modules/rds-aurora/aws"
  version        = "9.5.0"
  name           = "finance-pg"
  engine         = "aurora-postgresql"
  engine_version = "16.2"
  instance_class = "db.r6g.2xlarge"
  instances      = { one = {}, two = {} }
  vpc_id               = module.vpc.vpc_id
  storage_encrypted    = true
  deletion_protection  = true
  backup_retention_period = 14
}""", "HCL")

    add_heading_styled(doc, "4-5. msk.tf (Amazon MSK)", level=3)
    add_code_block(doc, """resource "aws_msk_cluster" "kafka" {
  cluster_name           = "finance-msk"
  kafka_version          = "3.7.x"
  number_of_broker_nodes = 6
  broker_node_group_info {
    instance_type   = "kafka.m7g.large"
    client_subnets  = module.vpc.private_subnets
    storage_info { ebs_storage_info { volume_size = 1000 } }
  }
  client_authentication { sasl { iam = true } }
  encryption_info {
    encryption_in_transit { client_broker = "TLS", in_cluster = true }
  }
}""", "HCL")

    # --- 2.5 CI/CD ---
    add_heading_styled(doc, "5. GitHub Actions — CI/CD 워크플로우", level=2)

    add_paragraph(doc, "build-test → containerize → deploy 3단계 파이프라인으로, OIDC를 통해 AWS 자격증명을 안전하게 획득합니다.",
                  font_size=10, space_after=6)

    add_table_with_style(doc,
        headers=["단계", "도구", "주요 작업"],
        rows=[
            ["Build & Test", "Maven + JDK 21", "빌드, 단위/통합 테스트, SonarQube 분석, CycloneDX SBOM"],
            ["Containerize", "Docker + ECR", "이미지 빌드, ECR 푸시, Trivy 보안 스캔"],
            ["Deploy", "Helm + Argo CD", "EKS 배포, --atomic 롤백, 5분 타임아웃"]
        ])

    # --- 2.6 Helm ---
    add_heading_styled(doc, "6. Helm 차트 — HPA / PDB / NetworkPolicy", level=2)

    add_heading_styled(doc, "6-1. HPA (Horizontal Pod Autoscaler)", level=3)
    add_table_with_style(doc,
        headers=["설정", "값", "설명"],
        rows=[
            ["minReplicas", "3", "최소 Pod 수"],
            ["maxReplicas", "20", "최대 Pod 수"],
            ["CPU Target", "65%", "CPU 사용률 기준"],
            ["Memory Target", "75%", "메모리 사용률 기준"],
            ["Custom Metric", "200 req/s", "HTTP 요청 기반 스케일링"],
            ["Scale Up Window", "30초", "급격한 확장 허용"],
            ["Scale Down Window", "300초", "안정화 후 축소"]
        ])

    add_heading_styled(doc, "6-2. PDB (Pod Disruption Budget)", level=3)
    add_paragraph(doc, "노드 업그레이드나 장애 발생 시 최소 2개의 Pod는 항상 서비스 가능 상태를 유지합니다. (minAvailable: 2)",
                  font_size=10, space_after=6)

    add_heading_styled(doc, "6-3. NetworkPolicy (Zero-Trust)", level=3)
    add_table_with_style(doc,
        headers=["방향", "허용 대상", "포트", "용도"],
        rows=[
            ["Ingress", "api-gateway Pod", "TCP 8080", "API Gateway에서만 접근 허용"],
            ["Egress", "kube-system namespace", "UDP 53", "DNS 조회"],
            ["Egress", "10.40.0.0/16", "TCP 5432", "PostgreSQL 접속"],
            ["Egress", "10.40.0.0/16", "TCP 9098", "Kafka MSK 접속"]
        ])

    add_heading_styled(doc, "운영 체크리스트", level=3)
    add_bullet_list(doc, [
        "점진 전환 단계: shadow read → canary 5% → 25% → 50% → 100%",
        "데이터 일관성: COMP-3 → BigDecimal 변환은 RoundingMode.HALF_EVEN으로 통일",
        "관측성: OpenTelemetry → Amazon Managed Grafana, 로그는 CloudWatch + Loki, 알람은 PagerDuty",
        "보안: RDS 및 MSK 모두 VPC 사설 서브넷, IAM IRSA로 시크릿 최소화, ECR 이미지 Trivy 게이트",
        "롤백: Helm --atomic + kubectl rollout undo, DB는 Flyway down 스크립트로 복구"
    ])

    add_page_break(doc)

    # ═══════════════════════════════════════════
    # PART 3: FINAL MIGRATION PLAN
    # ═══════════════════════════════════════════
    add_heading_styled(doc, "3부 — 최종 마이그레이션 계획서", level=1)

    add_info_box(doc, "본 계획서는 두 선행 리포트(AS-IS 시스템 분석 + TO-BE 마이그레이션 코드 패키지)를 통합하여 작성된 의사결정용 문서입니다. 한국 금융권 차세대 사업의 실전 표준을 기준으로 18개월 일정·예산·인력·리스크 전반을 다룹니다.")

    # --- 3.1 Executive Summary ---
    add_heading_styled(doc, "1. Executive Summary (경영진용 요약)", level=2)

    add_heading_styled(doc, "1-1. 배경", level=3)
    add_paragraph(doc, "삼성SDS가 운영하는 금융 차세대(Anyframe/F-ERP) 메인프레임은 안정적이지만, 연 18% 이상 증가하는 라이선스 비용, 메인프레임 인력의 평균 연령 54세, 신규 디지털 채널과의 통합 지연이라는 구조적 한계에 도달했습니다.",
                  font_size=10, space_after=6)

    add_heading_styled(doc, "1-2. 목표", level=3)
    add_table_with_style(doc,
        headers=["목표", "지표", "현재", "목표치"],
        rows=[
            ["TCO 절감", "5년 총비용", "1,262억 원", "1,059억 원 (▽16%)"],
            ["출시 리드타임 단축", "서비스 배포 주기", "12주", "3주 (75% 단축)"],
            ["시스템 가용성", "연간 가용률", "99.97%", "99.99%"],
            ["데이터 분석/AI", "실시간 스트림", "미지원", "Kafka 기반 실시간"]
        ])

    add_heading_styled(doc, "1-3. 기대효과", level=3)
    add_bullet_list(doc, [
        "신규 채널(오픈뱅킹·마이데이터) 통합 비용 60% 절감",
        "야간 배치 윈도우 단축 (현 6시간 → 2시간)",
        "신규 디지털 상품 출시 가속",
        "클라우드 네이티브 인재 채용 경쟁력 회복"
    ])

    add_heading_styled(doc, "1-4. 투자 규모", level=3)
    add_table_with_style(doc,
        headers=["Phase", "기간", "예산", "비중"],
        rows=[
            ["Phase 1: 기반 구축", "M1~M6", "75억 원", "31%"],
            ["Phase 2: 점진 전환", "M7~M12", "95억 원", "40%"],
            ["Phase 3: 완전 전환", "M13~M18", "70억 원", "29%"],
            ["합계", "18개월", "240억 원", "100%"]
        ])

    add_paragraph(doc, "비용 구성: 인건비 65%, 라이선스·인프라 20%, 외부 컨설팅 10%, 교육 5%",
                  font_size=9, color=RGBColor(0x66, 0x66, 0x66), space_after=6)

    add_heading_styled(doc, "1-5. 핵심 의사결정 포인트", level=3)
    add_numbered_list(doc, [
        "CIO/CTO 승인 — 18개월 동안 듀얼 런(레거시 + 신규) 운영비 부담을 감수할 의지",
        "금감원 사전 협의 — 전자금융감독규정 8조(시스템 변경) 및 클라우드 가이드라인 준수",
        "인력 구조조정 정책 결정 — COBOL 인력 재교육 vs 자연 감소 vs 전환 배치 비율",
        "벤더 락인 정책 — AWS 단일 vs 멀티클라우드 옵션"
    ])

    # --- 3.2 AS-IS vs TO-BE ---
    add_heading_styled(doc, "2. AS-IS vs TO-BE 아키텍처 비교", level=2)

    add_table_with_style(doc,
        headers=["영역", "AS-IS (메인프레임)", "TO-BE (AWS 클라우드 네이티브)"],
        rows=[
            ["컴퓨팅", "IBM z16, COBOL/PL/I, CICS, JCL Batch", "AWS EKS 1.30, Spring Boot 3.x (Java 21), Argo Workflows"],
            ["데이터베이스", "DB2 z/OS, IMS-DB, VSAM", "Aurora PostgreSQL 16, DynamoDB, S3 + Iceberg"],
            ["메시징", "MQ Series, FEP 고정전문", "Amazon MSK (Kafka 3.7), AWS API Gateway"],
            ["개발 환경", "RD/z, ChangeMan, 폐쇄망", "GitHub Enterprise, Backstage IDP, Codespaces"],
            ["CI/CD", "수동 빌드 + 변경관리 (평균 6주)", "GitHub Actions + Helm + Argo CD (평균 3일)"],
            ["관측성", "RMF, OMEGAMON", "OpenTelemetry → Prometheus/Grafana, CloudWatch, Loki"],
            ["보안", "RACF, 메인프레임 SAF", "IAM IRSA, AWS WAF, Vault, OPA Gatekeeper, KMS"],
            ["배포 단위", "Region-단위 LPAR, 전체 재기동", "Pod 단위 Rolling, Canary, BlueGreen"],
            ["DR (RTO/RPO)", "RTO 2h / RPO 5m (수동 절체)", "RTO 15m / RPO 30s (Multi-AZ + Cross-Region)"],
            ["확장성", "수직 확장 중심, 자원 사전 예약", "HPA/Karpenter 기반 자동 수평 확장"],
            ["인력 모델", "시스템 프로그래머·DBA·운영 분리", "DevSecOps Squad, SRE 4-on-call"],
            ["연 운영비(추정)", "약 220억 원", "약 145억 원 (5년차 안정화 기준)"]
        ])

    # --- 3.3 Roadmap ---
    add_heading_styled(doc, "3. 3단계 로드맵 (18개월)", level=2)

    add_heading_styled(doc, "Phase 1 — 기반 구축 (M1~M6)", level=3)
    add_paragraph(doc, "목표: 클라우드 플랫폼·CI/CD·관측성·CDC 파이프라인 완비, 1개 비핵심 서비스 시범 전환.",
                  font_size=10, bold=True, space_after=6)

    add_table_with_style(doc,
        headers=["마일스톤", "산출물", "인력"],
        rows=[
            ["M1: 거버넌스 수립", "클라우드 거버넌스 위원회, 금감원 사전협의서", "PMO 4 + 컴플 2"],
            ["M2: 랜딩존 구축", "AWS Control Tower, VPC, IAM, KMS, MSK 베이스", "인프라 6"],
            ["M3: 플랫폼 엔지니어링", "EKS 1.30, Backstage IDP, Argo CD, OPA, Vault", "플랫폼 8"],
            ["M4: 데이터 파이프라인", "IIDR/Debezium → MSK, EBCDIC→UTF-8 변환 라이브러리", "데이터 6"],
            ["M5: 시범 서비스 전환", "고객정보 조회(CST1) Strangler 배포, Shadow 운영", "앱 12"],
            ["M6: 운영 모델 전환", "SRE 온콜, 장애 RCA 프로세스, 카오스 엔지니어링", "SRE 5"]
        ])
    add_paragraph(doc, "투입 인력: 약 65명/월  |  예산: 75억 원", font_size=9, bold=True, color=SAMSUNG_BLUE, space_after=8)

    add_heading_styled(doc, "Phase 2 — 점진 전환 (M7~M12)", level=3)
    add_paragraph(doc, "목표: 핵심 트랜잭션 4종(BAL1·TRF1·LON1·CST1) 100% 전환, 대량 배치 30% 전환.",
                  font_size=10, bold=True, space_after=6)

    add_table_with_style(doc,
        headers=["마일스톤", "산출물", "인력"],
        rows=[
            ["M7: 잔액조회(BAL1) 전환", "Spring Boot 마이크로서비스, Canary 25%", "앱 12 + DBA 4"],
            ["M8: 이체(TRF1) 전환", "2PC 대체 Saga 패턴, Idempotency Key", "앱 14"],
            ["M9: 대출승인(LON1) 전환", "신용평가 모듈 + 룰엔진(Drools)", "앱 14 + DA 4"],
            ["M10: 배치 1차 전환", "Spring Batch + Argo Workflows, EOD 일부", "배치 10"],
            ["M11: 데이터 동기화 안정화", "양방향 CDC, DLQ 모니터링, 검증 자동화", "데이터 8"],
            ["M12: 컴플라이언스 점검", "금감원 보고, 중간 외부감사, 보안 모의해킹", "컴플 4 + 보안 6"]
        ])
    add_paragraph(doc, "투입 인력: 약 80명/월  |  예산: 95억 원", font_size=9, bold=True, color=SAMSUNG_BLUE, space_after=8)

    add_heading_styled(doc, "Phase 3 — 완전 전환 및 메인프레임 폐기 (M13~M18)", level=3)
    add_paragraph(doc, "목표: 모든 온라인·배치 100% 전환, 메인프레임 Read-Only 12주 운영 후 폐기, 인력 재배치 완료.",
                  font_size=10, bold=True, space_after=6)

    add_table_with_style(doc,
        headers=["마일스톤", "산출물", "인력"],
        rows=[
            ["M13: 잔존 배치·리포트 전환", "EOD/EOM 100%, 호스트 리포트 → PDF/Web", "배치 8 + 앱 8"],
            ["M14: 대외기관 전문 전환", "FEP → API Gateway + EAI, 표준전문 변환", "앱 6 + 대외 4"],
            ["M15: 데이터 마이그레이션 컷오버", "Big Bang Cutover Window 36h, Rollback 검증", "전사 통합팀"],
            ["M16: Read-Only 운영기", "메인프레임 12주 RO 모드, 거래 차이 모니터링", "운영 10"],
            ["M17: 메인프레임 폐기", "LPAR 종료, 라이선스 해지, Hardware 회수", "인프라 4"],
            ["M18: 안정화·인수인계", "운영 인수인계, 운영 매뉴얼 v2.0, 회고", "PMO 4 + SRE 6"]
        ])
    add_paragraph(doc, "투입 인력: 약 70명/월  |  예산: 70억 원", font_size=9, bold=True, color=SAMSUNG_BLUE, space_after=8)

    # --- 3.4 Risk Matrix ---
    add_heading_styled(doc, "4. 리스크 매트릭스 (5x5)", level=2)

    add_paragraph(doc, "확률(P) x 영향도(I) 기준 — 빨강(15이상), 주황(8~14), 노랑(4~7), 녹색(3이하)",
                  font_size=9, color=RGBColor(0x66, 0x66, 0x66), space_after=6)

    risk_table = doc.add_table(rows=13, cols=7)
    risk_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(risk_table, color=SAMSUNG_BLUE_HEX, sz="4")

    risk_headers = ["ID", "리스크", "P", "I", "PxI", "등급", "완화 방안"]
    risk_data = [
        ["R01", "COMP-3 → BigDecimal 변환 오차로 잔액 불일치", "4", "5", "20", "빨강", "변환 라이브러리 자동 회귀 테스트, 일자별 합계 대사 자동화"],
        ["R02", "컷오버 36h 윈도우 내 데이터 마이그레이션 실패", "3", "5", "15", "빨강", "모의 컷오버 3회, Rollback 스크립트, Read-Only 단계 운영"],
        ["R03", "금감원 컴플라이언스 지연", "3", "5", "15", "빨강", "사전협의 M1 착수, 외부 법무 자문 상시"],
        ["R04", "핵심 COBOL 인력 이탈", "4", "4", "16", "빨강", "핵심 인력 리텐션 보너스, 업무 매뉴얼 디지털화"],
        ["R05", "IMS-DB 계층형→관계형 전환 시 성능 저하", "4", "3", "12", "주황", "DynamoDB 채택, Aurora 인덱싱, Load Test"],
        ["R06", "대외 전문(EBCDIC) 인코딩 불일치", "3", "4", "12", "주황", "전문 매핑 사전 자동화, 30일 병행 운영"],
        ["R07", "클라우드 비용 통제 실패", "3", "4", "12", "주황", "FinOps 팀 신설, Budget Alert, Savings Plan"],
        ["R08", "카프카 메시지 순서·중복 처리 결함", "3", "3", "9", "주황", "Idempotency Key, 키별 파티셔닝, Outbox 패턴"],
        ["R09", "DR 사이트 동기화 지연", "2", "4", "8", "주황", "Multi-Region Active-Passive, RPO 30s 검증"],
        ["R10", "운영 인력의 클라우드 숙련도 부족", "4", "3", "12", "주황", "6개월 부트캠프, 외부 SRE 파트너십"],
        ["R11", "외부 SaaS(MSK, Aurora) 장애", "2", "4", "8", "주황", "Multi-AZ, 자체 Standby DB, AWS Enterprise Support"],
        ["R12", "보안 사고(Secret 유출, 공급망 공격)", "2", "5", "10", "주황", "IRSA + Vault + Trivy + SLSA Level 3"]
    ]

    # Header row
    for i, h in enumerate(risk_headers):
        cell = risk_table.rows[0].cells[i]
        cell.text = ""
        set_cell_shading(cell, SAMSUNG_BLUE_HEX)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=9, bold=True, color=WHITE)

    # Data rows with color-coded grade column
    grade_colors = {
        "빨강": "FFE0E0",
        "주황": "FFF0D0",
        "노랑": "FFFFD0",
        "녹색": "E0FFE0"
    }

    for r_idx, row_data in enumerate(risk_data):
        for c_idx, cell_text in enumerate(row_data):
            cell = risk_table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            if c_idx in [0, 2, 3, 4, 5]:  # Center-align ID, P, I, PxI, Grade
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text)
            font_size = 8 if c_idx in [1, 6] else 9
            set_font(run, size=font_size, color=DARK_GRAY)
            # Color grade cell
            if c_idx == 5 and cell_text in grade_colors:
                set_cell_shading(cell, grade_colors[cell_text])
                run.font.bold = True

    # --- 3.5 TCO ---
    add_heading_styled(doc, "5. TCO 비교 (5년, 단위: 억 원)", level=2)

    add_heading_styled(doc, "5-1. 메인프레임 유지 시나리오", level=3)
    add_table_with_style(doc,
        headers=["비용 항목", "Y1", "Y2", "Y3", "Y4", "Y5", "합계"],
        rows=[
            ["하드웨어 리스 (z16 LPAR)", "35", "35", "38", "41", "45", "194"],
            ["SW 라이선스(IBM, ISV)", "75", "82", "90", "99", "109", "455"],
            ["운영 인력(메인프레임 60명)", "72", "75", "78", "82", "86", "393"],
            ["전력·시설·DR", "18", "19", "20", "21", "22", "100"],
            ["외부 유지보수", "20", "22", "24", "26", "28", "120"],
            ["합계", "220", "233", "250", "269", "290", "1,262"]
        ])

    add_heading_styled(doc, "5-2. 클라우드 전환 시나리오", level=3)
    add_table_with_style(doc,
        headers=["비용 항목", "Y1(전환)", "Y2(전환)", "Y3", "Y4", "Y5", "합계"],
        rows=[
            ["일회성 전환 비용(SI/컨설팅/교육)", "150", "90", "0", "0", "0", "240"],
            ["AWS 인프라 (EKS·RDS·MSK 등)", "30", "50", "65", "70", "75", "290"],
            ["SaaS·옵저버빌리티", "8", "12", "14", "15", "16", "65"],
            ["운영 인력(SRE+개발 35명)", "50", "55", "60", "63", "66", "294"],
            ["듀얼 런(메인프레임 잔존)", "100", "60", "10", "0", "0", "170"],
            ["합계", "338", "267", "149", "148", "157", "1,059"]
        ])

    add_heading_styled(doc, "5-3. 차이 분석", level=3)
    add_bullet_list(doc, [
        "5년 누적 절감: 1,262 - 1,059 = 약 203억 원 (보수적 추정)",
        "일회성 전환 비용 회수 시점: 3.4년 차",
        "Y4부터 연 약 121억 원 절감, 신규 서비스 매출 기여(연 60~120억 추정)는 별도"
    ])

    # --- 3.6 Personnel ---
    add_heading_styled(doc, "6. 인력 전환 계획", level=2)

    add_heading_styled(doc, "6-1. 대상 인원 (현 110명 → 목표 85명)", level=3)
    add_table_with_style(doc,
        headers=["직군", "현재", "18개월 후", "전환 경로"],
        rows=[
            ["COBOL/PL/I 개발", "45", "12", "25명 Java 재교육, 5명 운영 전환, 3명 자연감소"],
            ["메인프레임 시스템 프로그래머", "12", "4", "6명 SRE 재교육, 2명 자연감소"],
            ["메인프레임 DBA", "10", "6", "4명 PostgreSQL/DynamoDB 전환 교육"],
            ["배치 운영", "18", "8", "8명 Argo/Spring Batch 전환, 2명 자연감소"],
            ["신규 채용(클라우드 네이티브)", "0", "30", "Java/K8s/SRE 외부 채용"],
            ["PMO/컴플라이언스", "8", "8", "유지"],
            ["QA/SDET", "6", "12", "6명 자동화 테스트 엔지니어 채용"]
        ])

    add_heading_styled(doc, "6-2. 6개월 재교육 커리큘럼", level=3)
    add_table_with_style(doc,
        headers=["월", "모듈", "시간", "인증"],
        rows=[
            ["M1", "Java 21 핵심 (OOP·Stream·Records·Virtual Thread)", "80h", "OCP Java SE"],
            ["M2", "Spring Boot 3.x + JPA + Test (Testcontainers)", "80h", "Spring Pro"],
            ["M3", "DevOps 기초 (Git·GitHub Actions·Docker)", "60h", "GitHub Foundations"],
            ["M4", "Kubernetes & Helm", "80h", "CKA"],
            ["M5", "AWS 핵심 (EKS·RDS·IAM·MSK·S3)", "80h", "AWS SAA"],
            ["M6", "사내 Capstone (실제 마이그레이션 모듈 1건)", "100h", "사내 인증"]
        ])

    add_paragraph(doc, "외부 채용 비율: 신규 인력의 35% (주로 SRE/Platform Engineer). 핵심 인력에게는 18개월 리텐션 보너스(연봉의 25%)를 제공하며, 자격증 비용 및 교육 시간은 회사가 부담합니다.",
                  font_size=9, color=RGBColor(0x66, 0x66, 0x66), space_after=8)

    # --- 3.7 Go/No-Go ---
    add_heading_styled(doc, "7. Go/No-Go 체크리스트", level=2)

    add_paragraph(doc, "각 Phase 종료 시점에 운영위원회가 공식 게이트 결정. 항목별 5단계 평가(0~4점), 합산 80% 이상 시 Go.",
                  font_size=10, space_after=8)

    checklist_sections = [
        ("7-1. 기술 (Technical)", [
            "단위 테스트 커버리지 >= 80%, 통합 테스트 >= 60%",
            "부하 테스트 결과: P99 응답 <= 500ms (BAL1), <= 1.5s (TRF1)",
            "데이터 대사(reconciliation): 일자별 잔액 합계 차이 = 0",
            "DR 시나리오: RTO <= 15m, RPO <= 30s 검증 완료",
            "모의 컷오버 2회 이상 성공 (Phase 3 한정)"
        ]),
        ("7-2. 비즈니스 (Business)", [
            "단위 서비스의 KPI(거래량·실패율) 정상 범위 유지",
            "사업 부서(영업·고객지원) 사용성 점검 완료",
            "신규 채널 통합 일정 영향 없음",
            "마케팅·고객 커뮤니케이션 완료"
        ]),
        ("7-3. 리스크 (Risk)", [
            "R01~R12 리스크 모든 High 항목 mitigation 완료",
            "미해결 P1 결함 0건, P2 결함 5건 이하",
            "보안 모의침투 결과 Critical 0건",
            "카오스 엔지니어링 시나리오 3건 이상 통과"
        ]),
        ("7-4. 규제·컴플라이언스 (Regulatory)", [
            "금감원 사전협의·후속 보고 완료",
            "전자금융감독규정 점검표(34개 항목) 100% 충족",
            "클라우드 이용 가이드라인(2024 개정판) 충족 보고서 제출",
            "개인정보 영향평가(PIA) 완료",
            "외부 감사(IT 컴플라이언스) 의견 무리한 사항 없음"
        ]),
        ("7-5. 운영 (Operations)", [
            "SRE 온콜 로테이션 운영 4주 이상 안정",
            "런북·플레이북 v1.0 작성·승인",
            "24/7 모니터링 알림 false positive 10% 이하",
            "BC/DR 훈련 완료",
            "인력 전환 계획상 핵심 인력 85% 이상 잔존"
        ])
    ]

    for section_title, items in checklist_sections:
        add_heading_styled(doc, section_title, level=3)
        # Use a table for checklist items
        for item in items:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.left_indent = Cm(1.0)
            run = p.add_run(f"☐  {item}")
            set_font(run, size=9, color=DARK_GRAY)

    add_page_break(doc)

    # ═══════════════════════════════════════════
    # APPENDIX
    # ═══════════════════════════════════════════
    add_heading_styled(doc, "부록", level=1)

    add_heading_styled(doc, "부록 A — 의사결정 일정", level=2)
    add_table_with_style(doc,
        headers=["시점", "의사결정", "책임자"],
        rows=[
            ["M0", "계획 승인 및 예산 확보", "CIO"],
            ["M2", "클라우드 벤더 단일/멀티 결정", "CTO"],
            ["M6", "Phase 1 → 2 게이트", "운영위원회"],
            ["M9", "핵심 인력 리텐션 보너스 집행", "CHRO"],
            ["M12", "Phase 2 → 3 게이트", "운영위원회"],
            ["M14", "Big Bang Cutover 일자 확정", "CIO + CRO"],
            ["M18", "메인프레임 폐기 공식 선언", "CEO 보고"]
        ])

    add_heading_styled(doc, "부록 B — 외부 참조 표준", level=2)
    add_bullet_list(doc, [
        "금융감독원 「전자금융감독규정」 8조, 14조의2",
        "금융위 「클라우드 컴퓨팅 서비스 이용 가이드라인」(2024 개정)",
        "ISO 22301 (BCMS — Business Continuity Management System)",
        "NIST SP 800-204C (마이크로서비스 보안)",
        "AWS Well-Architected Financial Services Lens"
    ])

    add_heading_styled(doc, "부록 C — 용어 사전 (Glossary)", level=2)
    add_table_with_style(doc,
        headers=["용어", "설명"],
        rows=[
            ["COBOL", "Common Business-Oriented Language. 메인프레임에서 사용하는 사무용 프로그래밍 언어"],
            ["CICS", "Customer Information Control System. IBM 온라인 트랜잭션 처리 시스템"],
            ["JCL", "Job Control Language. 메인프레임 배치 작업 정의 언어"],
            ["VSAM", "Virtual Storage Access Method. IBM 메인프레임 파일 관리 시스템"],
            ["IMS-DB", "Information Management System Database. IBM 계층형 데이터베이스"],
            ["DB2", "IBM 관계형 데이터베이스 관리 시스템"],
            ["COMP-3", "Packed Decimal. COBOL의 압축 10진수 데이터 형식"],
            ["EBCDIC", "Extended Binary Coded Decimal Interchange Code. IBM 문자 인코딩"],
            ["EKS", "Elastic Kubernetes Service. AWS 관리형 쿠버네티스 서비스"],
            ["MSK", "Managed Streaming for Apache Kafka. AWS 관리형 카프카 서비스"],
            ["Strangler Fig", "레거시 시스템을 점진적으로 대체하는 마이그레이션 패턴"],
            ["CDC", "Change Data Capture. 데이터베이스 변경사항 실시간 캡처 기술"],
            ["HPA", "Horizontal Pod Autoscaler. 쿠버네티스 자동 수평 확장 기능"],
            ["PDB", "Pod Disruption Budget. 쿠버네티스 Pod 중단 관리 정책"],
            ["IRSA", "IAM Roles for Service Accounts. EKS Pod별 AWS IAM 역할 매핑"],
            ["TCO", "Total Cost of Ownership. 총 소유 비용"],
            ["SRE", "Site Reliability Engineering. 사이트 신뢰성 엔지니어링"],
            ["Flyway", "자바 기반 데이터베이스 마이그레이션 도구"],
            ["Helm", "쿠버네티스 패키지 관리자"]
        ])

    # ═══════════════════════════════════════════
    # SAVE
    # ═══════════════════════════════════════════
    doc.save(OUTPUT_DOCX)
    print(f"DOCX saved: {OUTPUT_DOCX}")
    return OUTPUT_DOCX


if __name__ == "__main__":
    build_report()
