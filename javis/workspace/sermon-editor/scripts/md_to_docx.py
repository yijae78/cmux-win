"""
설교 원고 MD → DOCX 변환 스크립트
style-guide.md 기준: 나눔명조 15pt, 박스 14pt, 장평 85%, 자간 -3%
"""
import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Paths
BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SERMON_DIR = os.path.join(BASE, "20260325-로마서강해21-With")
MD_PATH = os.path.join(SERMON_DIR, "원고.md")
OUT_PATH = os.path.join(SERMON_DIR, "output", "With-함께라는이름-로마서강해21.docx")
IMG_DIR = os.path.join(SERMON_DIR, "images")


def set_font(run, name="나눔명조", size=15, bold=False, color=None):
    """Set font properties on a run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    # Set East Asian font
    rpr = run._element.get_or_add_rPr()
    ea = OxmlElement('w:rFonts')
    ea.set(qn('w:eastAsia'), name)
    rpr.insert(0, ea)
    if color:
        run.font.color.rgb = color


def set_spacing(paragraph, space_after=Pt(4), space_before=Pt(0)):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_after = space_after
    pf.space_before = space_before
    pf.line_spacing = Pt(22)


def add_image_if_exists(doc, img_path, width_cm=14):
    """Add image to document if file exists."""
    if os.path.exists(img_path):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run()
            run.add_picture(img_path, width=Cm(width_cm))
            return True
        except Exception as e:
            print(f"  Warning: Could not add image {img_path}: {e}")
    return False


def process_inline_formatting(paragraph, text, font_name="나눔명조", font_size=15, default_color=None):
    """Process **bold** and other inline formatting in text."""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            set_font(run, font_name, font_size, bold=True, color=default_color)
        else:
            if part:
                run = paragraph.add_run(part)
                set_font(run, font_name, font_size, color=default_color)


def convert_md_to_docx(md_path, out_path):
    """Convert sermon MD to DOCX with style mapping."""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Set default page margins (A4)
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    in_code_block = False
    in_table = False
    table_rows = []
    i = 0

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Skip empty &nbsp; lines - add spacing
        if line.strip() == '&nbsp;':
            doc.add_paragraph('')
            i += 1
            continue

        # Code blocks (``` ... ```) → grey background box
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
                i += 1
                continue
            else:
                in_code_block = False
                # Add code block as formatted paragraph
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                set_spacing(p)
                # Add shading
                pPr = p._element.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'E8E8E8')
                pPr.append(shd)
                run = p.add_run(code_text)
                set_font(run, "나눔명조", 14)
                i += 1
                continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table rows (| ... |)
        if line.strip().startswith('|') and '|' in line.strip()[1:]:
            # Check if separator row
            if re.match(r'\|[\s\-:]+\|', line.strip()):
                i += 1
                continue
            cells = [c.strip() for c in line.strip().split('|')[1:-1]]
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(cells)
            # Check if next line is still table
            if i + 1 < len(lines) and lines[i + 1].strip().startswith('|'):
                i += 1
                continue
            else:
                # End of table - create it
                in_table = False
                if table_rows:
                    cols = max(len(r) for r in table_rows)
                    tbl = doc.add_table(rows=len(table_rows), cols=cols)
                    tbl.style = 'Table Grid'
                    for ri, row in enumerate(table_rows):
                        for ci, cell in enumerate(row):
                            if ci < cols:
                                cell_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell)
                                cell_text = cell_text.replace('📖 ', '')
                                tbl.rows[ri].cells[ci].text = cell_text
                                for p in tbl.rows[ri].cells[ci].paragraphs:
                                    for run in p.runs:
                                        set_font(run, "나눔명조", 13)
                    table_rows = []
                i += 1
                continue

        # --- horizontal rule
        if line.strip() == '---':
            # Add a thin line or just spacing
            p = doc.add_paragraph()
            set_spacing(p, space_after=Pt(8), space_before=Pt(8))
            i += 1
            continue

        # # H1 - Series banner (22-24pt, bold, center)
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_spacing(p, space_after=Pt(2))
            run = p.add_run(text)
            set_font(run, "나눔명조", 22, bold=True)
            i += 1
            continue

        # ## H2 - Section headers (15pt, bold, red) or date/title
        if line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_paragraph()
            set_spacing(p, space_after=Pt(6), space_before=Pt(12))
            if 'Ⅰ.' in text or 'II.' in text or 'III.' in text or '■' in text or '▶' in text:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(text)
                set_font(run, "나눔명조", 15, bold=True, color=RGBColor(0xFF, 0, 0))
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                set_font(run, "나눔명조", 15, bold=True)
            i += 1
            continue

        # ### H3 - Sub-topic titles (15pt, bold, blue, underline)
        if line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_paragraph()
            set_spacing(p, space_after=Pt(6), space_before=Pt(10))
            run = p.add_run(text)
            set_font(run, "나눔명조", 15, bold=True, color=RGBColor(0, 0, 0xFF))
            run.font.underline = True
            i += 1
            continue

        # > blockquote
        if line.startswith('> '):
            text = line[2:].strip()
            p = doc.add_paragraph()
            set_spacing(p)
            pf = p.paragraph_format
            pf.left_indent = Cm(1)
            process_inline_formatting(p, text, "나눔명조", 15)
            i += 1
            continue

        # Image: ![alt](path)
        img_match = re.match(r'!\[([^\]]*)\]\(<([^>]+)>\)', line.strip())
        if not img_match:
            img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line.strip())
        if img_match:
            alt = img_match.group(1)
            img_file = img_match.group(2)
            img_path = os.path.join(SERMON_DIR, img_file)
            add_image_if_exists(doc, img_path, width_cm=13)
            i += 1
            continue

        # ▶ PPT directive lines (red, bold)
        if line.strip().startswith('▶'):
            text = line.strip()
            p = doc.add_paragraph()
            set_spacing(p)
            run = p.add_run(text)
            set_font(run, "나눔명조", 15, bold=True, color=RGBColor(0xFF, 0, 0))
            i += 1
            continue

        # * italic lines (hymn info etc) - 12pt
        if line.strip().startswith('*') and not line.strip().startswith('**') and not line.strip().startswith('*('):
            text = line.strip().lstrip('*').strip()
            p = doc.add_paragraph()
            set_spacing(p)
            run = p.add_run(text)
            set_font(run, "나눔명조", 12)
            run.font.italic = True
            i += 1
            continue

        # : definition lines (원어 분석)
        if line.strip().startswith(': '):
            text = line.strip()[2:]
            p = doc.add_paragraph()
            set_spacing(p)
            pf = p.paragraph_format
            pf.left_indent = Cm(0.5)
            process_inline_formatting(p, text, "나눔명조", 14)
            i += 1
            continue

        # (헬), (NIV), (NASB) lines
        if re.match(r'^\((헬|NIV|NASB|IT)\)', line.strip()):
            text = line.strip()
            p = doc.add_paragraph()
            set_spacing(p)
            pf = p.paragraph_format
            pf.left_indent = Cm(0.5)
            run = p.add_run(text)
            set_font(run, "나눔명조", 14)
            i += 1
            continue

        # Regular paragraph text
        text = line.strip()
        if text:
            p = doc.add_paragraph()
            set_spacing(p)
            process_inline_formatting(p, text, "나눔명조", 15)
        else:
            # Empty line
            pass

        i += 1

    # Save
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f"✓ DOCX saved: {out_path}")
    return out_path


if __name__ == '__main__':
    convert_md_to_docx(MD_PATH, OUT_PATH)
