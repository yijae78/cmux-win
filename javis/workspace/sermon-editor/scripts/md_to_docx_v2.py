#!/usr/bin/env python3
"""설교 원고 MD → DOCX 변환기 v2 (2026-04-22).

업그레이드 포인트 (대 md_to_docx.py):
  * CLI 인수 지원 (--input / --output / --clean-v2)
  * 헤더 색상 HTML 미리보기와 정렬 (H1 파랑 / H2·H3 빨강)
  * 볼드 강조 옅은 황금 (#D4AF37)
  * 성경 인용구 박스 (회색 배경 + 파란 좌측 테두리 + 14pt)
  * 축도 박스 (진한 파랑 테두리, 구분)
  * ▶PP{n} 슬라이드 마커 → 단독 황금 라벨 박스
  * --- 수평선 → 실제 가로선
  * 이미지 경로 MD 파일 기준 상대/절대 모두 지원, 한글 파일명 지원
  * 장평 85% / 자간 -3% 모든 본문 run에 적용
  * --clean-v2 플래그: 🔧 마커와 "수정 이력" 섹션 자동 제거 (설교 당일 최종본용)

사용법:
    python scripts/md_to_docx_v2.py --input <md파일>
    python scripts/md_to_docx_v2.py --input <md> --output <docx>
    python scripts/md_to_docx_v2.py --input <md> --clean-v2

결정 사항 정본: docs/docx_upgrade_decisions.md
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

# scripts 폴더를 sys.path 앞에 추가 (docx_style 하위 모듈 로드용)
THIS_DIR = Path(__file__).resolve().parent
if str(THIS_DIR) not in sys.path:
    sys.path.insert(0, str(THIS_DIR))

from docx_style import style as S  # noqa: E402
from docx_style import oxml_helpers as H  # noqa: E402


# ============================================================
# 공통 Run 스타일 적용
# ============================================================
def style_run(run, *, size_pt=None, bold=False, italic=False,
              color=None, underline=False, font=None,
              scale=True, spacing=True):
    """기본 run 스타일 묶음."""
    if font is None:
        font = S.FONT_BODY
    if size_pt is None:
        size_pt = S.SIZE_BODY
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if color is not None:
        run.font.color.rgb = color
    H.set_run_korean_font(run, font)
    if scale:
        H.set_run_scale(run, S.CHAR_WIDTH_SCALE)
    if spacing:
        H.set_run_spacing(run, S.CHAR_SPACING_TWIPS)


# ============================================================
# 문단 간격
# ============================================================
def set_para_spacing(para, *, before_pt=None, after_pt=None, line_pt=None):
    before_pt = S.PARA_SPACE_BEFORE_PT if before_pt is None else before_pt
    after_pt = S.PARA_SPACE_AFTER_PT if after_pt is None else after_pt
    line_pt = S.LINE_SPACING_PT if line_pt is None else line_pt
    pf = para.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    pf.line_spacing = Pt(line_pt)


# ============================================================
# 인라인 서식 처리 (**볼드** 등)
# ============================================================
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")


def add_inline(paragraph, text, *, size_pt=None, color=None, italic=False):
    """텍스트 내 **볼드** 처리하여 paragraph에 run들 추가."""
    size_pt = S.SIZE_BODY if size_pt is None else size_pt
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            style_run(run, size_pt=size_pt, italic=italic, color=color)
        bold_run = paragraph.add_run(m.group(1))
        style_run(bold_run, size_pt=size_pt, bold=True, italic=italic,
                  color=S.COLOR_BOLD)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        style_run(run, size_pt=size_pt, italic=italic, color=color)


# ============================================================
# 이미지 경로 해석 (한글 파일명 포함)
# ============================================================
def resolve_image_path(src, md_dir):
    """MD 파일 기준 상대 경로 / 절대 경로 모두 처리."""
    p = Path(src)
    if p.is_absolute() and p.exists():
        return p
    resolved = (md_dir / src).resolve()
    if resolved.exists():
        return resolved
    return None


def add_image_block(doc, img_src, md_dir):
    """이미지 삽입. 실패 시 회색 대체 박스."""
    path = resolve_image_path(img_src, md_dir)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if path and path.exists():
        run = p.add_run()
        try:
            run.add_picture(str(path), width=Cm(S.IMAGE_WIDTH_CM))
            return True
        except Exception as e:
            run.text = f"[이미지 삽입 실패: {path.name} ({e})]"
            style_run(run, size_pt=11, italic=True,
                      color=RGBColor(0x99, 0x99, 0x99))
            return False
    run = p.add_run(f"[이미지 없음: {img_src}]")
    style_run(run, size_pt=11, italic=True, color=RGBColor(0x99, 0x99, 0x99))
    return False


# ============================================================
# 박스 (성경 인용 / 축도 / PP 마커)
# ============================================================
def add_box(doc, lines_text, *, kind="scripture"):
    """1×1 표로 박스 생성.

    kind: 'scripture' | 'benediction' | 'pp' | 'monologue' (v3.1) | 'proclamation' (v3.1)
    """
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]

    # 배경 + 테두리 결정
    if kind == "benediction":
        bg = S.HEX_BENED_BG
        H.set_cell_shading(cell, bg)
        H.set_cell_border(cell, sides=["left"], hex_color=S.HEX_BENED_BORDER,
                          size_pt=2.0)
        H.set_cell_border(cell, sides=["top", "right", "bottom"],
                          hex_color=S.HEX_BENED_BORDER, size_pt=0.75)
    elif kind == "pp":
        bg = S.HEX_PP_BG
        H.set_cell_shading(cell, bg)
        H.set_cell_border(cell, hex_color=S.HEX_PP_BORDER, size_pt=1.0)
    elif kind == "monologue":
        # v3.1 신규 — 하나님 1인칭 선언 (금테 + 옅은 금배경 + 큰 폰트)
        bg = S.HEX_MONO_BG
        H.set_cell_shading(cell, bg)
        H.set_cell_border(cell, hex_color=S.HEX_MONO_BORDER, size_pt=2.0)
    elif kind == "proclamation":
        # v3.1 신규 — 회중 합송 선포문 (빨강 테 + 흰배경 + 중앙정렬 굵게)
        bg = S.HEX_PROC_BG
        H.set_cell_shading(cell, bg)
        H.set_cell_border(cell, hex_color=S.HEX_PROC_BORDER, size_pt=2.0)
    else:  # scripture
        bg = S.HEX_QUOTE_BG
        H.set_cell_shading(cell, bg)
        H.set_cell_border(cell, sides=["left"], hex_color=S.HEX_QUOTE_BORDER,
                          size_pt=2.5)
        # 다른 테두리는 옅게
        H.set_cell_border(cell, sides=["top", "right", "bottom"],
                          hex_color="E2E8F0", size_pt=0.5)

    # 셀 문단 설정
    first_para = cell.paragraphs[0]
    first = True
    for text_line in lines_text:
        if first:
            para = first_para
            first = False
        else:
            para = cell.add_paragraph()
        if kind == "pp":
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_para_spacing(para, before_pt=0, after_pt=0, line_pt=18)
            run = para.add_run(text_line)
            style_run(run, size_pt=S.SIZE_PP, bold=True, color=S.COLOR_PP_TEXT)
        elif kind == "benediction":
            set_para_spacing(para, before_pt=1, after_pt=1, line_pt=20)
            add_inline(para, text_line, size_pt=S.SIZE_BENEDICTION,
                       color=S.COLOR_BENED_TEXT)
        elif kind == "monologue":
            # 16pt 큰 폰트 + 좌측 들여쓰기
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.left_indent = Cm(0.5)
            set_para_spacing(para, before_pt=2, after_pt=2, line_pt=24)
            add_inline(para, text_line, size_pt=S.SIZE_MONOLOGUE,
                       color=S.COLOR_MONO_TEXT)
        elif kind == "proclamation":
            # 15pt 굵게 + 중앙정렬
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_para_spacing(para, before_pt=2, after_pt=2, line_pt=22)
            run = para.add_run(text_line)
            style_run(run, size_pt=S.SIZE_PROCLAMATION, bold=True,
                      color=S.COLOR_PROC_TEXT)
        else:
            set_para_spacing(para, before_pt=1, after_pt=1, line_pt=20)
            add_inline(para, text_line, size_pt=S.SIZE_QUOTE, italic=True)

    # 박스 뒤 여백
    trailer = doc.add_paragraph()
    set_para_spacing(trailer, before_pt=0, after_pt=4, line_pt=8)


# ============================================================
# 헤더
# ============================================================
HYMN_LINE_RE = re.compile(r"^\*[^*]")  # *로 시작하고 **(볼드) 아님
PP_MARKER_RE = re.compile(r"^▶PP(\d+)\s*$")
IMAGE_RE_BRACKET = re.compile(r"^!\[([^\]]*)\]\(<([^>]+)>\)$")
IMAGE_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)$")
TABLE_ROW_RE = re.compile(r"^\|.*\|\s*$")
TABLE_SEP_RE = re.compile(r"^\|[\s\-:|]+\|\s*$")
V2_MARKER_RE = re.compile(r"^>\s*🔧\s*\*\*\[v2")
V2_LOG_HEADER_RE = re.compile(r"^##\s+📝\s*수정\s*이력")
V2_TITLE_SUFFIX_RE = re.compile(r"\s*\*\*\(v2\)\*\*\s*$")
# v3.1 신규 — 박스 종류 자동 인식 (직전 라인이 호흡 지시문이면 박스 종류 변경)
MONO_TRIGGER_RE = re.compile(r"^\(IT:\s*(?:부드럽게|깊게)")
PROC_TRIGGER_RE = re.compile(r"^\(회중과 함께|^\(IT:.*?선포하듯")


def render_h1(doc, text, *, page_break=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if page_break:
        p.paragraph_format.page_break_before = True
    set_para_spacing(p, before_pt=S.H1_SPACE_BEFORE_PT,
                     after_pt=S.H1_SPACE_AFTER_PT)
    H.add_paragraph_border_bottom(p, hex_color=S.HEX_H1_UNDERLINE, size_pt=1.5)
    run = p.add_run(text)
    style_run(run, size_pt=S.SIZE_H1, bold=True, color=S.COLOR_H1)


def render_h2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p, before_pt=S.H2_SPACE_BEFORE_PT,
                     after_pt=S.H2_SPACE_AFTER_PT)
    run = p.add_run(text)
    style_run(run, size_pt=S.SIZE_H2, bold=True, color=S.COLOR_H2)


def render_h3(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p, before_pt=S.H3_SPACE_BEFORE_PT,
                     after_pt=S.H3_SPACE_AFTER_PT)
    run = p.add_run(text)
    style_run(run, size_pt=S.SIZE_H3, bold=True, color=S.COLOR_H3,
              underline=True)


# ============================================================
# 메인 변환 함수
# ============================================================
def convert(md_path, out_path, *, clean_v2=False):
    md_path = Path(md_path).resolve()
    out_path = Path(out_path).resolve()
    md_dir = md_path.parent

    if not md_path.exists():
        raise FileNotFoundError(f"입력 파일 없음: {md_path}")

    with md_path.open("r", encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()

    # 페이지 설정
    for section in doc.sections:
        section.page_width = Cm(S.PAGE_WIDTH_CM)
        section.page_height = Cm(S.PAGE_HEIGHT_CM)
        section.top_margin = Cm(S.MARGIN_TOP_CM)
        section.bottom_margin = Cm(S.MARGIN_BOTTOM_CM)
        section.left_margin = Cm(S.MARGIN_LEFT_CM)
        section.right_margin = Cm(S.MARGIN_RIGHT_CM)

    # 상태 변수
    in_bened_section = False          # ## 축도 섹션 안
    in_v2_log = False                 # ## 📝 수정 이력 섹션 안 (clean_v2)
    in_code_block = False
    code_buffer = []
    next_code_box_kind = None         # v3.1 — 직전 호흡 지시문이 박스 종류 결정 (monologue/proclamation)

    i = 0
    N = len(lines)

    while i < N:
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        # ---------- 코드 블록 ----------
        if stripped.startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_buffer = []
            else:
                in_code_block = False
                # v3.1: 직전 호흡 지시문에 따라 박스 종류 분기
                box_kind = next_code_box_kind if next_code_box_kind else "scripture"
                add_box(doc, code_buffer, kind=box_kind)
                code_buffer = []
                next_code_box_kind = None  # 1회용 — 다음 코드블록은 기본 scripture
            i += 1
            continue
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # ---------- v3.1: 호흡 지시문이면 다음 코드블록 박스 종류 예약 ----------
        if MONO_TRIGGER_RE.match(stripped):
            next_code_box_kind = "monologue"
        elif PROC_TRIGGER_RE.match(stripped):
            next_code_box_kind = "proclamation"

        # ---------- clean_v2: 🔧 마커 라인 스킵 ----------
        if clean_v2 and V2_MARKER_RE.match(stripped):
            i += 1
            # 직후 빈 줄도 같이 스킵
            if i < N and not lines[i].strip():
                i += 1
            continue

        # ---------- clean_v2: 수정 이력 섹션 전체 스킵 ----------
        if V2_LOG_HEADER_RE.match(stripped):
            if clean_v2:
                in_v2_log = True
                i += 1
                continue
        if in_v2_log:
            # 다음 ## 또는 파일 끝까지 스킵
            if stripped.startswith("## "):
                in_v2_log = False
                # 계속 처리 (아래 로직으로)
            else:
                i += 1
                continue

        # ---------- 축도 섹션 진입/이탈 ----------
        # v3.1: ## 축도 + # 축도 둘 다 인식 (원고 양식 자유도)
        if stripped == "## 축도" or stripped == "# 축도":
            in_bened_section = True
            # H1/H2 양식 그대로 렌더
            if stripped.startswith("## "):
                render_h2(doc, "축도")
            else:
                render_h1(doc, "축도")
            i += 1
            continue
        if in_bened_section and (stripped.startswith("## ") or stripped.startswith("# ")):
            in_bened_section = False
            # 아래 H1/H2 처리로 진행

        # ---------- 제목 라인 (## 본문 : ... (v2)) ----------
        if stripped.startswith("## 본문") and "제목" in stripped:
            text = stripped[3:]
            if clean_v2:
                text = V2_TITLE_SUFFIX_RE.sub("", text).rstrip()
            render_h2(doc, text)
            i += 1
            continue

        # ---------- 헤더 ----------
        if stripped.startswith("### "):
            render_h3(doc, stripped[4:])
            i += 1
            continue
        if stripped.startswith("## "):
            render_h2(doc, stripped[3:])
            i += 1
            continue
        if stripped.startswith("# "):
            h1_text = stripped[2:]
            is_sermon_start = "서론" in h1_text
            render_h1(doc, h1_text, page_break=is_sermon_start)
            i += 1
            continue

        # ---------- 수평선 ----------
        if stripped == "---" or stripped == "***":
            H.add_horizontal_rule(doc, hex_color=S.HEX_HR)
            i += 1
            continue

        # ---------- PP 슬라이드 마커 ----------
        m = PP_MARKER_RE.match(stripped)
        if m:
            num = m.group(1)
            add_box(doc, [f"▶ PP{num}"], kind="pp")
            i += 1
            continue

        # ---------- 이미지 ----------
        img_m = IMAGE_RE_BRACKET.match(stripped) or IMAGE_RE.match(stripped)
        if img_m:
            src = img_m.group(2)
            add_image_block(doc, src, md_dir)
            i += 1
            continue

        # ---------- 인용구 (>) 그룹핑 ----------
        if stripped.startswith(">"):
            quote_lines = []
            while i < N and lines[i].strip().startswith(">"):
                qstripped = lines[i].strip()
                # '> 🔧 [v2-X...]' 마커는 clean_v2 아니더라도 박스 안에 넣을지?
                # → 넣되 별도 표시 없이 인용으로. clean_v2는 위에서 처리됨.
                content = qstripped[1:].lstrip()
                quote_lines.append(content if content else "")
                i += 1
            # 빈 줄 연결된 경우 이어지는 > 있는지 확인
            # 축도 섹션이면 축도 박스, 아니면 성경 박스
            kind = "benediction" if in_bened_section else "scripture"
            add_box(doc, quote_lines, kind=kind)
            continue

        # ---------- 표 ----------
        if TABLE_ROW_RE.match(stripped) and not TABLE_SEP_RE.match(stripped):
            rows = []
            while i < N and TABLE_ROW_RE.match(lines[i].strip()):
                rs = lines[i].strip()
                if TABLE_SEP_RE.match(rs):
                    i += 1
                    continue
                cells = [c.strip() for c in rs.strip("|").split("|")]
                rows.append(cells)
                i += 1
            if rows:
                cols = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=cols)
                tbl.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci in range(cols):
                        cell_text = row[ci] if ci < len(row) else ""
                        # 인라인 서식 처리
                        cell_obj = tbl.rows[ri].cells[ci]
                        cell_obj.text = ""
                        cpara = cell_obj.paragraphs[0]
                        set_para_spacing(cpara, before_pt=1, after_pt=1,
                                         line_pt=18)
                        # 헤더 행은 굵게 + 배경
                        if ri == 0:
                            H.set_cell_shading(cell_obj, "E2E8F0")
                        # 볼드 마크다운 제거 후 처리
                        add_inline(cpara, cell_text, size_pt=13)
            continue

        # ---------- 찬송 정보 (*) ----------
        if HYMN_LINE_RE.match(stripped):
            text = stripped.lstrip("*").strip()
            p = doc.add_paragraph()
            set_para_spacing(p)
            run = p.add_run(text)
            style_run(run, size_pt=S.SIZE_HYMN_INFO, italic=True)
            i += 1
            continue

        # ---------- 빈 줄 ----------
        if not stripped:
            p = doc.add_paragraph()
            set_para_spacing(p, before_pt=0, after_pt=0, line_pt=10)
            i += 1
            continue

        # ---------- 일반 문단 ----------
        p = doc.add_paragraph()
        set_para_spacing(p)
        add_inline(p, stripped, size_pt=S.SIZE_BODY)
        # v3.1: 호흡 지시문 라인 아닌 일반 단락 등장 → 박스 예약 리셋
        #       (호흡 지시문 ↔ ``` 사이에 다른 본문 끼어든 경우 잘못 발동 방지)
        if not (MONO_TRIGGER_RE.match(stripped) or PROC_TRIGGER_RE.match(stripped)):
            next_code_box_kind = None
        i += 1

    # 저장
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


# ============================================================
# CLI
# ============================================================
def default_output_path(md_path):
    md_path = Path(md_path)
    return md_path.parent / "output" / md_path.with_suffix(".docx").name


def main():
    parser = argparse.ArgumentParser(
        description="설교 원고 MD → DOCX 변환기 v2")
    parser.add_argument("--input", "-i", required=True,
                        help="입력 MD 파일 경로 (예: sermons/xxx/원고_v2.md)")
    parser.add_argument("--output", "-o", default=None,
                        help="출력 DOCX 경로 (기본: <md파일_폴더>/output/<이름>.docx)")
    parser.add_argument("--clean-v2", action="store_true",
                        help="v2 메타 정리 모드 — 🔧 마커와 수정 이력 섹션 자동 제거 (최종본용)")
    args = parser.parse_args()

    md_path = Path(args.input)
    out_path = Path(args.output) if args.output else default_output_path(md_path)

    print(f"[md_to_docx_v2] Input:  {md_path}")
    print(f"[md_to_docx_v2] Output: {out_path}")
    print(f"[md_to_docx_v2] clean_v2: {args.clean_v2}")

    result = convert(md_path, out_path, clean_v2=args.clean_v2)
    print(f"✓ DOCX saved: {result}")


if __name__ == "__main__":
    main()
