"""시나리오 1 코드 리뷰 종합 보고서 — Word 문서 생성"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path
import datetime

doc = Document()

# ── 페이지 설정 ──
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ── 스타일 설정 ──
style = doc.styles['Normal']
font = style.font
font.name = 'Pretendard'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.4

# 한글 폰트 설정
rFonts = style.element.rPr.rFonts if style.element.rPr is not None else None
if rFonts is None:
    from docx.oxml import OxmlElement
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Pretendard')
    rPr.append(rFonts)

# ── 색상 상수 ──
DARK_BG = RGBColor(0x0f, 0x17, 0x2a)
ACCENT_BLUE = RGBColor(0x3b, 0x82, 0xf6)
ACCENT_GREEN = RGBColor(0x22, 0xc5, 0x5e)
ACCENT_RED = RGBColor(0xef, 0x44, 0x44)
ACCENT_AMBER = RGBColor(0xf5, 0x9e, 0x0b)
TEXT_MUTED = RGBColor(0x64, 0x74, 0x8b)
WHITE = RGBColor(0xff, 0xff, 0xff)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK_BG if level <= 2 else RGBColor(0x1e, 0x29, 0x3b)
    return h

def add_bold_text(paragraph, text, color=None):
    run = paragraph.add_run(text)
    run.bold = True
    if color:
        run.font.color.rgb = color
    return run

def add_text(paragraph, text, color=None, bold=False, size=None):
    run = paragraph.add_run(text)
    if color:
        run.font.color.rgb = color
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    return run

def set_cell_shading(cell, color_hex):
    """셀 배경색 설정"""
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def make_table(headers, rows, col_widths=None, highlight_col=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '0f172a')

    # 데이터
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'f1f5f9')
            if highlight_col is not None and c_idx == highlight_col:
                # 점수 열 색상
                try:
                    score = float(str(val).split('/')[0])
                    if score >= 7:
                        run.font.color.rgb = ACCENT_GREEN
                    elif score >= 5.5:
                        run.font.color.rgb = ACCENT_AMBER
                    else:
                        run.font.color.rgb = ACCENT_RED
                    run.bold = True
                except:
                    pass

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

def add_severity_badge(paragraph, severity):
    colors = {'High': ACCENT_RED, 'Medium': ACCENT_AMBER, 'Low': ACCENT_GREEN, 'Critical': ACCENT_RED}
    color = colors.get(severity, TEXT_MUTED)
    run = paragraph.add_run(f' [{severity}] ')
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(9)

# ════════════════════════════════════════
# 문서 본문 시작
# ════════════════════════════════════════

# ── 표지 ──
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(title, 'Javis Fleet', color=ACCENT_BLUE, bold=True, size=14)
add_text(title, ' 시나리오 1', color=DARK_BG, bold=True, size=14)

main_title = doc.add_paragraph()
main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(main_title, '코드 리뷰 종합 보고서', color=DARK_BG, bold=True, size=28)

doc.add_paragraph('')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(subtitle, 'EnvironmentScan-system-main-v4', color=TEXT_MUTED, size=12)

doc.add_paragraph('')
doc.add_paragraph('')

# 메타 정보
meta_items = [
    ('대상', 'EnvironmentScan-system-main-v4 (환경 스캔 시스템)'),
    ('파일', 'monitor.py (745줄), launch_monitor.py (37줄)'),
    ('분석일', '2026-06-17'),
    ('참여 워커', 'Worker1(Claude), Worker2(AGY), Worker3(Codex)'),
    ('작성', 'Javis Master (Claude Code CLI)'),
]
for label, value in meta_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, f'{label}: ', color=TEXT_MUTED, bold=True, size=10)
    add_text(p, value, color=DARK_BG, size=10)

doc.add_page_break()

# ── 목차 ──
add_heading_styled('목차', level=1)
toc_items = [
    '1. Executive Summary',
    '2. 3개 워커 공통 지적 사항 (Critical)',
    '3. 2개 워커 이상 지적 사항 (High)',
    '4. 개별 워커 고유 발견 사항',
    '5. 주요 버그 목록',
    '6. 개선 우선순위 로드맵',
    '7. 워커별 원본 보고서',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.runs[0].font.color.rgb = ACCENT_BLUE

doc.add_page_break()

# ── 1. Executive Summary ──
add_heading_styled('1. Executive Summary', level=1)

p = doc.add_paragraph()
add_text(p, '3개 워커의 독립적 분석을 종합한 결과, EnvironmentScan 대시보드는 ')
add_text(p, '프로토타입으로서 기능적으로 완성', bold=True, color=ACCENT_GREEN)
add_text(p, '되어 있으나, 프로덕션 수준의 ')
add_text(p, '보안·성능·유지보수성', bold=True, color=ACCENT_RED)
add_text(p, '에서 개선이 필요합니다.')

doc.add_paragraph('')

make_table(
    ['평가 항목', '점수', '요약'],
    [
        ['기능 완성도', '7/10', '실시간 ETA, SVG 차트, 자동 갱신 등 핵심 기능 구현 완료'],
        ['코드 품질', '6/10', '단일 파일 모놀리스, 200줄+ 함수, 변수명 개선 필요'],
        ['보안', '5/10', 'XSS 위험, 네트워크 노출, 입력값 미검증'],
        ['성능', '5.5/10', '매 렌더링마다 파일시스템 스캔, blocking sleep'],
        ['에러 핸들링', '4.5/10', '포괄적 예외 묵살, 디버깅 불가'],
        ['아키텍처', '5.5/10', '데이터/로직/렌더링 미분리, 테스트 불가 구조'],
        ['종합', '5.6/10', '프로토타입 OK, 프로덕션 배포 전 리팩토링 필수'],
    ],
    col_widths=[3.5, 2, 10.5],
    highlight_col=1
)

doc.add_paragraph('')

# 종합 점수 강조
score_p = doc.add_paragraph()
score_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(score_p, '종합 점수: ', color=TEXT_MUTED, bold=True, size=14)
add_text(score_p, '5.6 / 10', color=ACCENT_AMBER, bold=True, size=24)

doc.add_page_break()

# ── 2. 공통 지적 사항 ──
add_heading_styled('2. 3개 워커 공통 지적 사항 (Critical — 전원 일치)', level=1)

# 2.1
add_heading_styled('2.1 XSS/HTML 인젝션 위험', level=2)
items = [
    ('Worker1', 'CSS/HTML 문자열에 외부 데이터 직접 삽입 (unsafe_allow_html=True)'),
    ('Worker2', 'JSON 파일 값이 HTML에 포매팅되어 직접 렌더링 — XSS 공격 가능'),
    ('Worker3', 'html.escape() 없이 signal titles, category labels 등 삽입'),
]
for worker, desc in items:
    p = doc.add_paragraph()
    add_text(p, f'• {worker}: ', bold=True, color=ACCENT_BLUE)
    add_text(p, desc)

rec = doc.add_paragraph()
add_text(rec, '▶ 권장: ', bold=True, color=ACCENT_GREEN)
add_text(rec, '모든 데이터 파생 값에 html.escape(str(value)) 적용')

doc.add_paragraph('')

# 2.2
add_heading_styled('2.2 time.sleep() Blocking 문제', level=2)
items = [
    ('Worker1', 'time.sleep(30) + st.rerun() — UI 무응답 유발'),
    ('Worker2', '메인 스레드 차단형 대기로 반응성 저해'),
    ('Worker3', '세션 워커를 5~30초 점유 → 다중 사용자 시 심각'),
]
for worker, desc in items:
    p = doc.add_paragraph()
    add_text(p, f'• {worker}: ', bold=True, color=ACCENT_BLUE)
    add_text(p, desc)

rec = doc.add_paragraph()
add_text(rec, '▶ 권장: ', bold=True, color=ACCENT_GREEN)
add_text(rec, 'st_autorefresh 또는 타임스탬프 기반 rerun으로 교체')

doc.add_paragraph('')

# 2.3
add_heading_styled('2.3 포괄적 예외 묵살 (except Exception: return None)', level=2)
items = [
    ('Worker1', '_j() 함수의 모든 에러 묵살 → 디버깅 불가'),
    ('Worker2', 'try-except-pass 구문으로 원인 추적 불가'),
    ('Worker3', 'FileNotFoundError, JSONDecodeError 등 구분 없이 동일 처리'),
]
for worker, desc in items:
    p = doc.add_paragraph()
    add_text(p, f'• {worker}: ', bold=True, color=ACCENT_BLUE)
    add_text(p, desc)

rec = doc.add_paragraph()
add_text(rec, '▶ 권장: ', bold=True, color=ACCENT_GREEN)
add_text(rec, '좁은 범위 예외 캐치 + 로깅 추가')

doc.add_page_break()

# ── 3. High 지적 사항 ──
add_heading_styled('3. 2개 워커 이상 지적 사항 (High)', level=1)

add_heading_styled('3.1 매 렌더링마다 파일시스템 전체 스캔', level=2)
for worker, desc in [
    ('Worker2', 'st.session_state / st.cache_data 미활용 → 불필요한 디스크 I/O'),
    ('Worker3', '파일 glob + stat 반복 → 로그 축적 시 성능 저하'),
]:
    p = doc.add_paragraph()
    add_text(p, f'• {worker}: ', bold=True, color=ACCENT_BLUE)
    add_text(p, desc)
rec = doc.add_paragraph()
add_text(rec, '▶ 권장: ', bold=True, color=ACCENT_GREEN)
add_text(rec, 'st.cache_data로 (path, mtime) 기반 캐싱')

doc.add_paragraph('')

add_heading_styled('3.2 단일 파일 모놀리스 구조', level=2)
for worker, desc in [
    ('Worker1', '데이터 로딩, 비즈니스 로직, 렌더링이 모두 한 파일에 혼재'),
    ('Worker3', '데이터 접근/상태 도출/렌더링 3계층 분리 권장'),
]:
    p = doc.add_paragraph()
    add_text(p, f'• {worker}: ', bold=True, color=ACCENT_BLUE)
    add_text(p, desc)
rec = doc.add_paragraph()
add_text(rec, '▶ 권장: ', bold=True, color=ACCENT_GREEN)
add_text(rec, 'data_loader.py / view_model.py / renderer.py 분리')

doc.add_paragraph('')

add_heading_styled('3.3 타입 안전성 부족', level=2)
for worker, desc in [
    ('Worker1', '_j() → dict | None 반환, 내부 구조 미검증'),
    ('Worker3', 'TypedDict 또는 dataclass로 JSON 스키마 명시 필요'),
]:
    p = doc.add_paragraph()
    add_text(p, f'• {worker}: ', bold=True, color=ACCENT_BLUE)
    add_text(p, desc)

doc.add_page_break()

# ── 4. 개별 워커 고유 발견 사항 ──
add_heading_styled('4. 개별 워커 고유 발견 사항', level=1)

add_heading_styled('Worker1 (Claude) 고유', level=2)
for item in [
    'CSS 키프레임 중복 정의 (CSS vs CSS_ANIM 상수)',
    'Three Horizons 데이터 로딩 시 today vs data_date 불일치',
    'launch_monitor의 stdout/stderr 소거로 디버깅 불가',
    '종합 점수: 6.0/10',
]:
    p = doc.add_paragraph(style='List Bullet')
    p.text = item

doc.add_paragraph('')

add_heading_styled('Worker2 (AGY — Gemini) 고유', level=2)
for item in [
    'window.parent.document 접근 — iframe 샌드박스 우회 우려 (ZOOM 스크립트)',
    '명시적 바인딩 주소 부재 → 네트워크 노출 위험',
    'Streamlit 상태 관리 아키텍처(st.session_state) 미준수',
    '강한 파일시스템 종속성과 도메인 레이어 결합',
]:
    p = doc.add_paragraph(style='List Bullet')
    p.text = item

doc.add_paragraph('')

add_heading_styled('Worker3 (Codex — OpenAI) 고유', level=2)
for item in [
    '_load() 내부 parsed dict 직접 변이(mutation) → 캐싱 시 사이드이펙트',
    '날짜 처리 시 historical vs current-day 데이터 혼합 가능',
    '런처의 고정 포트(8504) 충돌 가능성',
    'KeyboardInterrupt 시 자식 프로세스 미종료',
]:
    p = doc.add_paragraph(style='List Bullet')
    p.text = item

doc.add_page_break()

# ── 5. 주요 버그 목록 ──
add_heading_styled('5. 주요 버그 목록', level=1)

bugs = [
    ['1', 'High', 'monitor.py', 'XSS — JSON 값이 HTML에 미이스케이프 삽입'],
    ['2', 'High', 'monitor.py', 'time.sleep(30) blocking → UI 무응답'],
    ['3', 'High', 'monitor.py', '_j() 포괄 예외 묵살 → 오류 원인 추적 불가'],
    ['4', 'Medium', 'monitor.py', 'Three Horizons today vs data_date 불일치'],
    ['5', 'Medium', 'monitor.py', 'CSS 키프레임 중복 정의'],
    ['6', 'Medium', 'monitor.py', '매 렌더링마다 파일시스템 glob+stat 반복'],
    ['7', 'Medium', 'launch_monitor.py', 'stdout/stderr 소거 → 실패 시 디버깅 불가'],
    ['8', 'Medium', 'launch_monitor.py', '고정 포트 충돌 + 준비 상태 미확인'],
    ['9', 'Low', 'monitor.py', 'wf_validation, cross_wf 로드 후 미사용'],
    ['10', 'Low', 'launch_monitor.py', '프로세스 종료 시 자식 프로세스 누수'],
]

bug_table = doc.add_table(rows=1 + len(bugs), cols=4)
bug_table.alignment = WD_TABLE_ALIGNMENT.CENTER
bug_table.style = 'Table Grid'

for i, h in enumerate(['#', '심각도', '파일', '내용']):
    cell = bug_table.rows[0].cells[i]
    cell.text = ''
    run = cell.paragraphs[0].add_run(h)
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = WHITE
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(cell, '0f172a')

for r_idx, row in enumerate(bugs):
    for c_idx, val in enumerate(row):
        cell = bug_table.rows[r_idx + 1].cells[c_idx]
        cell.text = ''
        run = cell.paragraphs[0].add_run(str(val))
        run.font.size = Pt(9)
        if r_idx % 2 == 1:
            set_cell_shading(cell, 'f1f5f9')
        # 심각도 컬러
        if c_idx == 1:
            severity_colors = {'High': ACCENT_RED, 'Medium': ACCENT_AMBER, 'Low': ACCENT_GREEN}
            run.font.color.rgb = severity_colors.get(val, TEXT_MUTED)
            run.bold = True

doc.add_page_break()

# ── 6. 개선 우선순위 ──
add_heading_styled('6. 개선 우선순위 로드맵', level=1)

add_heading_styled('Phase 1 — 즉시 (보안·안정성)', level=2)
for num, item in [
    ('1', 'html.escape() 적용 (XSS 차단)'),
    ('2', 'time.sleep() → st_autorefresh 교체'),
    ('3', '_j() 예외 범위 축소 + 로깅'),
]:
    p = doc.add_paragraph()
    add_text(p, f'  {num}. ', bold=True, color=ACCENT_RED)
    add_text(p, item)

doc.add_paragraph('')

add_heading_styled('Phase 2 — 단기 (구조·성능)', level=2)
for num, item in [
    ('4', 'st.cache_data 도입 (파일 I/O 감소)'),
    ('5', '_load() 함수 분할 (175줄 → 3~4개 함수)'),
    ('6', 'Three Horizons 날짜 정합성 수정'),
]:
    p = doc.add_paragraph()
    add_text(p, f'  {num}. ', bold=True, color=ACCENT_AMBER)
    add_text(p, item)

doc.add_paragraph('')

add_heading_styled('Phase 3 — 중기 (아키텍처)', level=2)
for num, item in [
    ('7', '3계층 분리 (data / view_model / render)'),
    ('8', 'TypedDict 도입으로 타입 안전성 확보'),
    ('9', '런처 개선 (준비 상태 폴링, 에러 출력)'),
]:
    p = doc.add_paragraph()
    add_text(p, f'  {num}. ', bold=True, color=ACCENT_GREEN)
    add_text(p, item)

doc.add_page_break()

# ── 7. 워커별 원본 보고서 ──
add_heading_styled('7. 워커별 원본 보고서', level=1)

make_table(
    ['워커', '파일', '크기'],
    [
        ['Worker1 (Claude)', 'worker1_code_analysis.md', '11,807 bytes'],
        ['Worker2 (AGY)', 'worker2_security_review.md', '7,114 bytes'],
        ['Worker3 (Codex)', 'worker3_tech_review.md', '8,798 bytes'],
    ],
    col_widths=[4, 6, 3]
)

doc.add_paragraph('')
doc.add_paragraph('')

# 푸터
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(footer, '— ', color=TEXT_MUTED, size=9)
add_text(footer, 'Javis Fleet', color=ACCENT_BLUE, bold=True, size=9)
add_text(footer, ' 시나리오 1 테스트 산출물 —', color=TEXT_MUTED, size=9)

doc.add_paragraph('')
footer2 = doc.add_paragraph()
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(footer2, '마스터(Claude Code CLI)가 3개 워커의 독립 분석을 종합하여 작성', color=TEXT_MUTED, size=9)

footer3 = doc.add_paragraph()
footer3.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(footer3, f'생성일: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M")}', color=TEXT_MUTED, size=8)

# ── 저장 ──
output_path = Path(r'C:\dev\cmux-win\javis\tests\output\scenario1\코드리뷰_종합보고서_시나리오1.docx')
doc.save(str(output_path))
print(f"Word 문서 생성 완료: {output_path}")
print(f"파일 크기: {output_path.stat().st_size:,} bytes")
