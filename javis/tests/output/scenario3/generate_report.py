"""시나리오 3 교차 프로젝트 분석 종합 보고서 Word 문서 생성."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path

doc = Document()

# ── 스타일 설정 ──
style = doc.styles["Normal"]
style.font.name = "맑은 고딕"
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

DARK = RGBColor(0x0C, 0x11, 0x1B)
ACCENT = RGBColor(0x3B, 0x82, 0xF6)
GREEN = RGBColor(0x22, 0xC5, 0x5E)
RED = RGBColor(0xEF, 0x44, 0x44)
AMBER = RGBColor(0xF5, 0x9E, 0x0B)
GRAY = RGBColor(0x64, 0x74, 0x8B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT if level == 1 else DARK
    return h


def add_body(text, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table


def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    return p


def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = ACCENT
    run.font.size = Pt(9)
    return p


# ════════════════════════════════════════════
# 표지
# ════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading("시나리오 3\n교차 프로젝트 분석\n종합 보고서", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = ACCENT
    run.font.size = Pt(28)

doc.add_paragraph()
meta_items = [
    ("대상", "Office-Monitor (사무실 영상 모니터링) × EnvironmentScan (환경 스캔 대시보드)"),
    ("분석일", "2026-06-17"),
    ("참여", "Worker1(Claude), Worker2(AGY), Worker3(Sub-agent), Worker4(Sub-agent), Worker5(AGY)"),
    ("생성", "Javis Fleet Master"),
]
for label, value in meta_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_l = p.add_run(f"{label}: ")
    run_l.bold = True
    run_l.font.color.rgb = GRAY
    run_v = p.add_run(value)
    run_v.font.color.rgb = DARK

doc.add_page_break()

# ════════════════════════════════════════════
# Executive Summary
# ════════════════════════════════════════════
add_heading("Executive Summary")
add_body(
    '두 프로젝트는 "모니터링"이라는 공통 도메인을 공유하지만, 근본적으로 다른 패러다임을 채택한다. '
    "Office-Monitor는 능동형 에지 디바이스 모니터링(YOLO+InsightFace+ByteTrack+OpenCV), "
    "EnvironmentScan은 수동형 파일 기반 대시보드(Streamlit+JSON 파싱)로 설계되었다."
)

add_table(
    ["구분", "Office-Monitor", "EnvironmentScan"],
    [
        ["유형", "능동형 에지 디바이스 모니터링", "수동형 파일 기반 대시보드"],
        ["핵심 기술", "YOLO + InsightFace + ByteTrack + OpenCV", "Streamlit + JSON 파싱"],
        ["데이터 흐름", "Push/Streaming (카메라 실시간)", "Pull/Polling (파일 주기적 읽기)"],
        ["UI", "PyQt6 네이티브 GUI", "Streamlit 웹 대시보드"],
        ["복잡도", "높음 (983줄 God Class, 멀티스레드+프로세스)", "중간 (746줄 단일 파일, 단일 스레드)"],
    ],
)

# ════════════════════════════════════════════
# 1. 프로젝트별 종합 점수
# ════════════════════════════════════════════
add_heading("1. 프로젝트별 종합 점수")

add_table(
    ["평가 항목", "Office-Monitor", "EnvironmentScan"],
    [
        ["아키텍처", "7/10", "6.5/10"],
        ["보안", "5/10", "7/10"],
        ["성능", "6.5/10", "8/10"],
        ["에러 처리", "6/10", "5.5/10"],
        ["종합", "6.1/10", "6.8/10"],
    ],
)

# ════════════════════════════════════════════
# 2. 교차 분석: 공통 발견사항
# ════════════════════════════════════════════
add_heading("2. 교차 분석: 공통 발견사항")

add_heading("2.1 공통 강점", level=2)
add_bullet("관심사 분리 의지: 두 프로젝트 모두 데이터 수집과 렌더링을 분리하려는 설계 의도가 명확")
add_bullet("상태 모델 존재: Office-Monitor는 track_id 기반 상태 추적, EnvironmentScan은 State dataclass로 3상태 모델 구현")
add_bullet("자동 복구 메커니즘: Office-Monitor의 카메라 재연결 지수 백오프, EnvironmentScan의 모드별 자동 갱신")

add_heading("2.2 공통 약점", level=2)
add_bullet("God Class 경향: DetectionThread 983줄, monitor.py 746줄 — 두 프로젝트 모두 단일 거대 모듈")
add_bullet('에러 삼킴 패턴: except Exception: pass 또는 except: pass 양쪽 모두 존재')
add_bullet("테스트 부재: 두 프로젝트 모두 단위 테스트 파일 없음")
add_bullet("설정 관리 미흡: 임계값, 경로 등이 하드코딩 또는 분산")

# ════════════════════════════════════════════
# 3. 보안 비교
# ════════════════════════════════════════════
add_heading("3. 보안 비교 (기술 검증 결과)")

add_heading("Office-Monitor — 중대 보안 이슈", level=2)
add_table(
    ["심각도", "항목", "설명"],
    [
        ["높음", "생체 정보 평문 저장", "얼굴 이미지 + 임베딩이 암호화 없이 저장. GDPR/개인정보보호법 위반 가능"],
        ["높음", "Queue 미정리 리소스 누수", "multiprocessing.Queue close()/join_thread() 미호출 → Windows 좀비 스레드"],
        ["중간", "모델 파일 무결성 미검증", "YOLO .pt 파일 pickle 역직렬화 공격 벡터"],
        ["중간", "유사도 임계값 0.4", "오인식(False Positive) 위험 — 보안 목적이면 0.5+ 권장"],
    ],
)

add_heading("EnvironmentScan — 경미한 보안 이슈", level=2)
add_table(
    ["심각도", "항목", "설명"],
    [
        ["중간", "CSS unsafe_allow_html", "XSS 벡터 가능성 (Streamlit 이미 격리하지만 원칙적 위험)"],
        ["낮음", "파일 경로 하드코딩", "설정 파일에서 읽지 않고 코드 내 직접 지정"],
    ],
)

# ════════════════════════════════════════════
# 4. 성능 비교
# ════════════════════════════════════════════
add_heading("4. 성능 비교")

add_heading("Office-Monitor 성능 병목", level=2)
add_bullet("_rebuild_matrix() — 임베딩 1개 변경 시 전체 행렬 재구성 (O(N))")
add_bullet("_is_duplicate_face() — O(N) 선형 탐색 매 프레임")
add_bullet("frame.copy() 과다 — 1080p 프레임 6MB x 다수 복사")
add_bullet("time.sleep() 기반 녹화 타이밍 — Windows 15ms 정밀도 한계")

add_heading("EnvironmentScan 성능 특성", level=2)
add_bullet("읽기 전용이므로 성능 병목 미미")
add_bullet("JSON 파싱이 주요 연산 — 현재 규모에서 문제 없음")
add_bullet("5초/30초 폴링 주기가 적절")

# ════════════════════════════════════════════
# 5. 모니터링 모범 패턴
# ════════════════════════════════════════════
add_heading("5. 모니터링 모범 패턴 (교차 추출)")

add_heading("5.1 Office-Monitor에서 추출한 우수 사례", level=2)
add_table(
    ["패턴", "설명", "적용 범위"],
    [
        ["GIL 우회 프로세스 분리", "CPU 바운드 추론을 multiprocessing.Process로 격리", "AI/ML 모니터링 전반"],
        ["최고 프레임 선별 (Best Frame)", "3초 버퍼 → 품질 점수 최고 1장만 저장", "영상/이미지 기반 감지"],
        ["카메라 지수 백오프 재연결", "1→2→4→8→16→30초 점진적 재시도", "하드웨어 연동 모니터링"],
        ["Poison Pill 종료", "None → join(5) → terminate() 3단계", "멀티프로세스 시스템"],
    ],
)

add_heading("5.2 EnvironmentScan에서 추출한 우수 사례", level=2)
add_table(
    ["패턴", "설명", "적용 범위"],
    [
        ["State Dataclass 중심 데이터 흐름", "로딩 → State 객체 → 렌더링 단방향", "대시보드/관측 시스템"],
        ["모드별 갱신 주기", "live=5초, completed=30초, idle=갱신없음", "리소스 효율적 폴링"],
        ["WF_ORDER 중앙 관리", "워크플로우 순서/라벨을 상수로 관리", "다단계 파이프라인 시각화"],
    ],
)

add_heading("5.3 교차 적용 제안", level=2)
add_bullet("Office-Monitor → EnvironmentScan 패턴 적용: State dataclass 도입으로 DetectionThread 상태 관리 개선")
add_bullet("EnvironmentScan → Office-Monitor 패턴 적용: 모드별 폴링 주기 개념을 카메라 프레임 레이트에 적용 (야간=저FPS, 감지 시=고FPS)")
add_bullet("공통 추출 가능 모듈: 로깅 설정, 에러 핸들러, 설정 로더")

# ════════════════════════════════════════════
# 6. 개선 우선순위
# ════════════════════════════════════════════
add_heading("6. 개선 우선순위 (종합)")

add_heading("Phase 1 — 즉시 (보안/안정성)", level=2)
add_bullet("1. Office-Monitor: 생체 정보 암호화 (AES + 접근 제어)")
add_bullet("2. Office-Monitor: mp.Queue.close() + join_thread() 추가")
add_bullet("3. Office-Monitor: 모델 파일 체크섬 검증")
add_bullet("4. 양쪽: except Exception: pass → logger.warning() 전환")

add_heading("Phase 2 — 단기 (성능)", level=2)
add_bullet("5. Office-Monitor: _rebuild_matrix() 증분 업데이트")
add_bullet("6. Office-Monitor: time.monotonic() 기반 녹화 타이밍")
add_bullet("7. Office-Monitor: 프레임 복사 최소화 (공유 메모리 고려)")

add_heading("Phase 3 — 중기 (아키텍처)", level=2)
add_bullet("8. Office-Monitor: DetectionThread God Class 분리 (FaceMatcher, CaptureManager, VisitLogger)")
add_bullet("9. EnvironmentScan: monitor.py 3분할 (state/render/app)")
add_bullet("10. 양쪽: 단위 테스트 도입")

# ════════════════════════════════════════════
# 7. 시나리오 3 테스트 검증 결과
# ════════════════════════════════════════════
add_heading("7. 시나리오 3 테스트 검증 결과")

add_table(
    ["검증 항목", "결과"],
    [
        ["8번째 패널(Worker5) 동적 생성", "성공 — %8 pane 생성"],
        ["8-pane 균등분할 적용", "성공 — workspace.set_layout API"],
        ["Worker5 라벨 'Worker5(AGY)'", "성공 — surface.rename API"],
        ["대시보드에 8개 에이전트 표시", "성공 — 7터미널 + 1대시보드"],
        ["5워커 동시 작업 수행", "부분 성공 — W1/W2/W5 정상, W4 지연, Codex 입력 실패"],
        ["교차 프로젝트 분석 결과", "성공 — 5개 산출물 생성 (sub-agent 2개 대체)"],
        ["RSI 산출물 생성", "성공 — 교차 프로젝트 보고서 + 모범 패턴 문서"],
    ],
)

add_heading("발견된 cmux-win 이슈", level=2)
add_bullet("Codex 프롬프트 입력 실패: tmux send-keys로 전송한 한국어 장문 프롬프트가 Codex CLI에서 처리 안 됨")
add_bullet("Worker4 장시간 지연: Claude 워커의 파일 권한 프롬프트 처리 지연 (--dangerously-skip-permissions 미적용)")
add_bullet("AGY 자동 시작 실패: tmux split-window -h \"agy\" 실행 시 셸만 열리고 AGY 미시작 → 수동 agy 명령 필요")

# ════════════════════════════════════════════
# 8. 워커별 원본 보고서
# ════════════════════════════════════════════
add_heading("8. 워커별 원본 보고서")

add_table(
    ["워커", "파일", "크기", "담당"],
    [
        ["Worker1(Claude)", "office_monitor_review.md", "15,316 bytes", "Office-Monitor 코드 리뷰"],
        ["Worker2(AGY)", "monitoring_patterns.md", "6,675 bytes", "모니터링 패턴 비교"],
        ["Sub-agent (Codex 대체)", "technical_verification.md", "19,619 bytes", "보안/성능 기술 검증"],
        ["Sub-agent (W4 대체)", "env_scan_review.md", "10,432 bytes", "EnvironmentScan 코드 리뷰"],
        ["Worker5(AGY)", "architecture_comparison.md", "11,262 bytes", "아키텍처 비교"],
    ],
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Javis Fleet 시나리오 3 교차 프로젝트 분석 — 마스터 종합 보고서")
run.italic = True
run.font.color.rgb = GRAY
run.font.size = Pt(9)

# ── 저장 ──
out = Path(r"C:\dev\cmux-win\javis\tests\output\scenario3\교차프로젝트분석_종합보고서_시나리오3.docx")
doc.save(str(out))
print(f"Saved: {out} ({out.stat().st_size:,} bytes)")
