"""시나리오 2 종합 보고서 Word 문서 생성."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path

doc = Document()

# ── 스타일 설정 ──
style = doc.styles["Normal"]
style.font.name = "맑은 고딕"
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

DARK = RGBColor(0x0C, 0x11, 0x1B)
ACCENT = RGBColor(0x3B, 0x82, 0xF6)
GREEN = RGBColor(0x22, 0xC5, 0x5E)
RED = RGBColor(0xEF, 0x44, 0x44)
AMBER = RGBColor(0xF5, 0x9E, 0x0B)
GRAY = RGBColor(0x64, 0x74, 0x8B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT if level == 1 else DARK
    return h


def add_body(text, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table


def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    return p


def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = ACCENT
    run.font.size = Pt(9)
    return p


# ════════════════════════════════════════════
# 표지
# ════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading("시나리오 2\nGlobalNews 아키텍처 분석\n종합 보고서", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = ACCENT
    run.font.size = Pt(28)

doc.add_paragraph()
meta_items = [
    ("대상", "GlobalNews-Crawling-AgenticWorkflow"),
    ("분석 모듈", "src/crawling/, src/analysis/, src/storage/, src/utils/"),
    ("분석일", "2026-06-17"),
    ("참여 워커", "Worker1(Claude), Worker2(AGY), Worker3(Codex), Worker4(Claude)"),
    ("생성", "Javis Fleet Master"),
]
for label, value in meta_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_l = p.add_run(f"{label}: ")
    run_l.bold = True
    run_l.font.color.rgb = GRAY
    run_v = p.add_run(value)
    run_v.font.color.rgb = DARK

doc.add_page_break()

# ════════════════════════════════════════════
# Executive Summary
# ════════════════════════════════════════════
add_heading("Executive Summary")
add_body("4개 워커의 독립적 분석을 종합한 결과, GlobalNews 시스템은 운영급(production-grade) 크롤링 인프라와 연구급(research-grade) 분석 파이프라인을 결합한 높은 완성도의 시스템입니다. 다만 두 모듈 간 인터페이스 계약과 상태 관리 일관성에서 개선이 필요합니다.")

add_table(
    ["평가 항목", "점수", "요약"],
    [
        ["크롤링 아키텍처", "7.5/10", "Conductor 패턴 + 4-Level 재시도 + Never-Abandon 우수"],
        ["분석 파이프라인", "7.5/10", "8단계 순차 + 체크포인트 + 메모리 관리 우수"],
        ["모듈 간 인터페이스", "6/10", "JSONL→Parquet 안정적, 메타데이터 유실"],
        ["보안", "7/10", "내부 경로 생성, TOCTOU/atomic write 불일치"],
        ["에러 처리", "8/10", "계층적 예외 + Circuit Breaker + Graceful Degradation"],
        ["전체 아키텍처", "7/10", "Staged Monolith 적절, 복구 상태 SOT 분산"],
        ["종합", "7.2/10", "인터페이스·상태 관리 정비 필요"],
    ],
)

# ════════════════════════════════════════════
# 1. 공통 지적 사항
# ════════════════════════════════════════════
add_heading("1. 4개 워커 공통 지적 사항 (Critical)")

add_heading("1.1 crawler.py와 pipeline.py의 역할 중복", level=2)
add_bullet("Worker1: Crawler 클래스가 CrawlingPipeline과 거의 동일한 파이프라인을 독립 구현 — Dead Code")
add_bullet("Worker3: pipeline.py가 주 경로, crawler.py는 Writer/State만 재사용")
add_bullet("Worker2: 크롤링 레이어 내부 오케스트레이터 중복 확인")
add_quote("권장: Crawler 클래스를 deprecated하고 CrawlingPipeline으로 통합")

add_heading("1.2 assert 문의 프로덕션 사용", level=2)
add_bullet("Worker1: pipeline.py 전반에 assert 남용 (L1108, 1113, 1358 등)")
add_bullet("Worker3: 런타임 필수 조건이 assert에 의존 — python -O에서 방어 소실")
add_bullet("Worker4: 분석 파이프라인에서도 동일 패턴")
add_quote("권장: assert → if ... is None: raise RuntimeError(...) 명시적 예외로 전환")

add_heading("1.3 Atomic Write 불일치", level=2)
add_bullet("Worker1: _save_bypass_state()에 atomic write 미적용")
add_bullet("Worker2: Parquet 쓰기도 최종 경로에 직접 기록")
add_bullet("Worker4: docstring에 temp+rename 명시했으나 실제 미구현")
add_quote("권장: 모든 상태/데이터 파일에 temp → os.replace() 패턴 통일")

# ════════════════════════════════════════════
# 2. 2개 워커 이상 지적
# ════════════════════════════════════════════
add_heading("2. 2개 워커 이상 지적 사항 (High)")

add_heading("2.1 크롤링-분석 메타데이터 유실", level=2)
add_bullet("Worker2(AGY): crawl_method, crawl_tier, is_paywall_truncated가 Parquet 변환 시 누락")
add_bullet("Worker3(Codex): 분석 모듈이 JSONL 텍스트만 읽어 장애 컨텍스트 인계 불가")
add_quote("권장: ARTICLES_PA_SCHEMA에 크롤링 메타 필드 추가")

add_heading("2.2 God Class — CrawlingPipeline (2,377줄)", level=2)
add_bullet("Worker1: URL 발견, 추출, 중복 제거, 재시도, bypass, 상태 영속화를 모두 담당")
add_bullet("Worker3: private field 직접 접근 (self._guard._circuit_breakers)")
add_quote("권장: NeverAbandonRunner, BypassHandler, FailureReporter 등 분리")

add_heading("2.3 메모리 유틸리티 중복", level=2)
add_bullet("Worker4: _get_rss_mb() vs MemoryMonitor.get_rss_gb() 중복")
add_bullet("Worker4: os.uname() 사용 — Windows에서 AttributeError")
add_quote("권장: MemoryMonitor 하나로 통일, Windows 호환 확보")

# ════════════════════════════════════════════
# 3. 개별 워커 고유 발견
# ════════════════════════════════════════════
add_heading("3. 개별 워커 고유 발견 사항")

add_heading("Worker1(Claude) — 크롤링 모듈", level=2)
add_bullet("RawArticle frozen dataclass 재생성 비용 → dataclasses.replace() 권장")
add_bullet("CrawlState.mark_site_complete()의 processed_urls 타입 불일치 (list vs set)")
add_bullet("_merge_result()의 articles.extend() — H-16 메모리 최적화와 모순")
add_bullet("RSS fallback 코드 약 40줄 중복")
add_bullet("compute_content_hash() 한국어 정규화 한계")

add_heading("Worker2(AGY) — 인터페이스", level=2)
add_bullet("RawArticle.source_id → Parquet source 축소 결합, source_name 누락")
add_bullet("SQLite unicode61 FTS5 토크나이저 융합 우수")
add_bullet("크롤러 .crawl_state.json과 분석 레이어 간 헬스 모니터링 동기화 부재")

add_heading("Worker3(Codex) — 아키텍처", level=2)
add_bullet("복구 상태 SOT 분산 (CrawlState, RetryManager, CircuitBreaker 등)")
add_bullet("self_recovery.py가 main.py 실행 경로에 미통합")
add_bullet("Circuit Breaker 계층 중복 (3곳)")
add_bullet("실패/성공 URL이 같은 processed set — resume semantics 약화")

add_heading("Worker4(Claude) — 분석 모듈", level=2)
add_bullet("Stage 1 단일 스레드 순차 처리 — 멀티코어 미활용")
add_bullet("langdetect 매 호출 시 seed 재설정 + detect()+detect_langs() 이중 호출")
add_bullet("Stage 3 zero-shot: emotion(8)+STEEPS(6) 이중 추론 → 통합 시 50% 절감")
add_bullet("_estimate_coverage() O(N^2) 코사인 유사도 병목")
add_bullet("Kiwi 모델 하드코딩 경로 C:/kiwi_model")

# ════════════════════════════════════════════
# 4. 모듈별 종합 점수
# ════════════════════════════════════════════
add_heading("4. 모듈별 종합 점수")
add_table(
    ["모듈", "아키텍처", "보안", "성능", "에러 처리", "종합"],
    [
        ["crawling", "7/10", "7/10", "7/10", "8/10", "7.3"],
        ["analysis", "8/10", "6/10", "6/10", "7/10", "6.8"],
        ["storage", "8/10", "7/10", "8/10", "7/10", "7.5"],
        ["interface", "6/10", "-", "-", "5/10", "5.5"],
    ],
)

# ════════════════════════════════════════════
# 5. 개선 우선순위
# ════════════════════════════════════════════
add_heading("5. 개선 우선순위")

add_heading("Phase 1 — 즉시 (데이터 무결성/안전성)", level=2)
add_bullet("1. CrawlState.mark_site_complete() 타입 불일치 수정 (list→set)")
add_bullet("2. Atomic write 패턴 전체 통일 (bypass_state, Parquet 출력)")
add_bullet("3. assert → 명시적 예외 전환")
add_bullet("4. Windows 호환: os.uname() → platform.system() 통일")

add_heading("Phase 2 — 단기 (성능/효율)", level=2)
add_bullet("5. Stage 1 병렬화 (ProcessPoolExecutor 또는 Kiwi pipe())")
add_bullet("6. Stage 3 zero-shot 통합 추론 (14 라벨 한 번에)")
add_bullet("7. langdetect 최적화 (seed 1회, detect_langs() 단독)")
add_bullet("8. dataclasses.replace() 적용")

add_heading("Phase 3 — 중기 (아키텍처)", level=2)
add_bullet("9. Crawler 클래스 deprecated → CrawlingPipeline 통합")
add_bullet("10. CrawlingPipeline God Class 분리 (2,377줄 → 4~5개 모듈)")
add_bullet("11. Parquet 스키마에 크롤링 메타데이터 추가")
add_bullet("12. 복구 상태 SOT 통합 (단일 상태 저장소)")
add_bullet("13. self_recovery.py main.py 실행 경로 통합")

# ════════════════════════════════════════════
# 6. 워커별 원본 보고서
# ════════════════════════════════════════════
add_heading("6. 워커별 원본 보고서")
add_table(
    ["워커", "파일", "크기", "담당"],
    [
        ["Worker1(Claude)", "crawling_analysis.md", "11,662 bytes", "크롤링 모듈 코드 리뷰"],
        ["Worker2(AGY)", "cross_module_review.md", "6,878 bytes", "모듈 간 인터페이스 리뷰"],
        ["Worker3(Codex)", "architecture_review.md", "20,103 bytes", "전체 아키텍처 리뷰"],
        ["Worker4(Claude)", "analysis_module_review.md", "13,602 bytes", "분석 모듈 코드 리뷰"],
    ],
)

# ════════════════════════════════════════════
# 7. 시나리오 2 테스트 검증
# ════════════════════════════════════════════
add_heading("7. 시나리오 2 테스트 검증 결과")
add_table(
    ["검증 항목", "결과"],
    [
        ["7번째 패널(Worker4) 동적 생성", "성공"],
        ["7-pane 균등분할 재조정", "성공"],
        ["Worker4 라벨 설정", "성공"],
        ["대시보드 Worker4 자동 반영", "성공 (제한 수정 후)"],
        ["Worker1과 Worker4 병렬 작업", "성공"],
        ["4개 워커 동시 작업중", "확인"],
        ["워커 간 결과물 파일 공유", "성공"],
        ["Claude 워커 컨텍스트 관리", "이슈 발견 (43k CLAUDE.md)"],
        ["최종 산출물 저장", "성공 (5개 파일)"],
    ],
)

doc.add_paragraph()
add_heading("발견된 cmux-win 이슈", level=2)
add_bullet("대시보드 터미널 제한 [:5]: Worker4 추가 시 표시 안 됨 → 제한 제거 + 동적 감지로 수정")
add_bullet("Claude 워커 컨텍스트 소진: 43k CLAUDE.md + 대용량 소스 = 빠른 컨텍스트 풀")

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Javis Fleet 시나리오 2 테스트 — 마스터 종합 보고서")
run.italic = True
run.font.color.rgb = GRAY
run.font.size = Pt(9)

# ── 저장 ──
out = Path(r"C:\dev\cmux-win\javis\tests\output\scenario2\아키텍처분석_종합보고서_시나리오2.docx")
doc.save(str(out))
print(f"Saved: {out} ({out.stat().st_size:,} bytes)")
